import gradio as gr
import requests
import json
import os
import tempfile
from pathlib import Path
import time
from typing import Generator, Tuple, List
import re
import subprocess

# 文档处理库
try:
    import PyPDF2
except ImportError:
    print("请安装PyPDF2: pip install PyPDF2")

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("请安装python-docx: pip install python-docx")

try:
    import pdfplumber
except ImportError:
    print("请安装pdfplumber: pip install pdfplumber")

class DocumentProcessor:
    """文档处理类"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """从PDF文件提取文本"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- 第 {page_num + 1} 页 ---\n"
                        text += page_text
        except Exception as e:
            print(f"PDF读取错误: {e}")
            # 备用方法
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text += f"\n--- 第 {page_num + 1} 页 ---\n"
                        text += page.extract_text()
            except Exception as e2:
                print(f"备用PDF读取也失败: {e2}")
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """从DOCX文件提取文本"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"DOCX读取错误: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """从TXT文件提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except Exception as e:
                print(f"TXT读取错误: {e}")
                return ""
        except Exception as e:
            print(f"TXT读取错误: {e}")
            return ""

class OllamaClient:
    """Ollama客户端"""
    
    def __init__(self, base_url: str = "http://localhost:11434"):
        self.base_url = base_url
        self.model = "gemma3:12b"
    
    def generate_stream(self, prompt: str, context: str = "") -> Generator[str, None, None]:
        """流式生成响应"""
        if context:
            full_prompt = f"{prompt}\n\n文本内容：\n{context}"
        else:
            full_prompt = prompt
        
        payload = {
            "model": self.model,
            "prompt": full_prompt,
            "stream": True,
            "options": {
                "temperature": 0.7,
                "num_ctx": 4096
            }
        }
        
        try:
            response = requests.post(
                f"{self.base_url}/api/generate",
                json=payload,
                stream=True,
                timeout=300
            )
            
            if response.status_code == 200:
                for line in response.iter_lines():
                    if line:
                        try:
                            data = json.loads(line.decode('utf-8'))
                            if 'response' in data:
                                yield data['response']
                            if data.get('done', False):
                                break
                        except json.JSONDecodeError:
                            continue
            else:
                yield f"错误: HTTP {response.status_code}"
                
        except requests.exceptions.RequestException as e:
            yield f"连接错误: {str(e)}"

class TextSplitter:
    """文本分割器"""
    
    @staticmethod
    def split_text(text: str, max_length: int = 3000) -> List[str]:
        """将文本分割成较小的块"""
        if len(text) <= max_length:
            return [text]
        
        # 按段落分割
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) <= max_length:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # 如果单个段落太长，按句子分割
        final_chunks = []
        for chunk in chunks:
            if len(chunk) <= max_length:
                final_chunks.append(chunk)
            else:
                sentences = re.split(r'[.!?]+', chunk)
                temp_chunk = ""
                for sentence in sentences:
                    if len(temp_chunk) + len(sentence) <= max_length:
                        temp_chunk += sentence + ". "
                    else:
                        if temp_chunk:
                            final_chunks.append(temp_chunk.strip())
                        temp_chunk = sentence + ". "
                if temp_chunk:
                    final_chunks.append(temp_chunk.strip())
        
        return final_chunks

class PromptEnhancer:
    """提示词优化器"""
    
    def __init__(self):
        self.ollama = OllamaClient()
        self.enhancement_methods = self._get_enhancement_methods()
    
    def _get_enhancement_methods(self):
        """获取不同的提示词优化方法"""
        return {
            "链式思维 (CoT)": """你是专业的提示词工程师。请将以下提示词优化为链式思维风格，适合19岁广州女大学生（商务英语二年级，对文化运动、旅游和AI感兴趣）使用。

原提示词：{original_prompt}

请直接给出优化后的中文提示词（不要只给建议），要求：
1. 使用"让我一步步思考..."的开头
2. 将任务分解为清晰的步骤
3. 语言亲切活泼，适合年轻女性
4. 适当融入文化、旅游、AI等兴趣点

优化后的提示词：""",

            "树状思维 (ToT)": """你是专业的提示词工程师。请将以下提示词优化为树状思维风格，适合19岁广州女大学生（商务英语二年级，对文化运动、旅游和AI感兴趣）使用。

原提示词：{original_prompt}

请直接给出优化后的中文提示词（不要只给建议），要求：
1. 使用"让我探索不同的方法..."开头
2. 引导AI生成多种解决方案并比较
3. 语言自然亲切，符合年轻女性表达习惯
4. 融入广州本地文化和兴趣元素

优化后的提示词：""",

            "图状思维 (GoT)": """你是专业的提示词工程师。请将以下提示词优化为图状思维风格，适合19岁广州女大学生（商务英语二年级，对文化运动、旅游和AI感兴趣）使用。

原提示词：{original_prompt}

请直接给出优化后的中文提示词（不要只给建议），要求：
1. 使用"让我梳理这些概念之间的联系..."开头
2. 强调概念间的关系和连接
3. 语言生动形象，便于理解复杂关系
4. 结合AI技术与商业应用的关联

优化后的提示词：""",

            "全维思维 (EoT)": """你是专业的提示词工程师。请将以下提示词优化为全维思维风格，适合19岁广州女大学生（商务英语二年级，对文化运动、旅游和AI感兴趣）使用。

原提示词：{original_prompt}

请直接给出优化后的中文提示词（不要只给建议），要求：
1. 使用"让我从各个角度全面分析..."开头
2. 从多个维度和角度分析问题
3. 语言全面细致，体现女性细腻思维
4. 结合广州国际化商业环境特色

优化后的提示词：""",

            "CO-STAR框架": """你是专业的提示词工程师。请将以下提示词优化为CO-STAR框架格式，适合19岁广州女大学生（商务英语二年级，对文化运动、旅游和AI感兴趣）使用。

原提示词：{original_prompt}

请直接给出优化后的中文提示词（不要只给建议），严格按照以下格式：

背景：[结合广州本地特色和大学生活的背景设定]
目标：[明确要完成的任务]  
风格：[适合年轻女大学生的输出风格]
语调：[温和、鼓励的情感色调]
受众：[19岁商务英语专业女学生]
回应：[期望的响应格式和结构]

[在此处给出完整的结构化提示词，要亲切自然有温度]

优化后的提示词："""
        }
    
    def enhance_prompt(self, original_prompt: str, method: str, progress_callback=None) -> str:
        """优化提示词"""
        if method not in self.enhancement_methods:
            return "未知的优化方法"
        
        if progress_callback:
            progress_callback(f"正在使用{method}优化提示词...")
        
        enhancement_prompt = self.enhancement_methods[method].format(original_prompt=original_prompt)
        
        result = ""
        for response_part in self.ollama.generate_stream(enhancement_prompt):
            result += response_part
            if progress_callback:
                progress_callback(f"生成中: {result[-50:]}")
        
        return result

class DocumentAnalyzer:
    """文档分析主类"""
    
    def __init__(self):
        self.processor = DocumentProcessor()
        self.ollama = OllamaClient()
        self.splitter = TextSplitter()
        
        # 定义基础提示词
        self.base_prompts = self._get_base_prompts()
        
        # 定义思维模式前缀
        self.thinking_modes = self._get_thinking_modes()
    
    def _get_thinking_modes(self):
        """获取不同思维模式的前缀指令"""
        return {
            "标准模式": "",
            
            "链式思维 (CoT)": """Before providing your final answer, think step by step:
1. First, identify the key information in the text
2. Then, analyze the relationships between different concepts
3. Next, consider the implications and connections
4. Finally, synthesize your findings into a coherent response

Let me work through this systematically:
""",
            
            "树状思维 (ToT)": """I'll explore multiple possible approaches to analyze this text, then choose the best path:

**Approach 1**: Focus on main themes and arguments
**Approach 2**: Analyze from historical/contextual perspective  
**Approach 3**: Examine practical applications and implications
**Approach 4**: Look at theoretical frameworks and concepts

Let me evaluate each approach and select the most comprehensive one:
""",
            
            "图状思维 (GoT)": """I'll analyze this text by mapping the interconnected relationships between ideas:

**Core Concepts** → **Supporting Evidence** → **Implications**
        ↓                    ↓                    ↓
**Related Themes** → **Counterarguments** → **Applications**
        ↓                    ↓                    ↓
**Historical Context** → **Current Relevance** → **Future Considerations**

Now let me trace these connections systematically:
""",
            
            "全维思维 (EoT)": """I'll analyze this from every possible angle and dimension:

🔍 **Analytical Dimensions:**
- Semantic analysis (what does it say?)
- Pragmatic analysis (what does it do?)
- Critical analysis (what are the strengths/weaknesses?)
- Contextual analysis (what's the broader picture?)
- Predictive analysis (what are the implications?)

🎯 **Perspective Angles:**
- Academic/scholarly view
- Practical/applied view  
- Critical/skeptical view
- Creative/innovative view
- Holistic/systemic view

Let me synthesize insights from all these dimensions:
"""
        }
    
    def _get_base_prompts(self):
        """获取基础提示词（不包含思维模式前缀）"""
        
        study_guide_prompt = """Create a comprehensive study guide for a chapter or section titled <CHAPTER_OR_SECTION_TITLE> based on the provided text. The study guide should include:

1. **Summary**
   - Provide a concise 200-word summary in English that captures the main points and key arguments of the text.
   - Focus on the most important concepts and their relationships.

2. **Comprehension Questions**  
   - A series of short-answer questions that focus on the chapter's key concepts.  
   - Each answer should require 2–3 sentences.

3. **Analytical Essay Questions**  
   - A set of open-ended prompts that invite critical evaluation and exploration of broader implications related to the chapter's themes.  
   - Encourage connections to real-world issues, ethical considerations, or theoretical debates.

4. **Glossary of Terms**  
   - A list of the main technical or thematic terms introduced in the chapter.  
   - Provide concise, student-friendly definitions for each term.

Ensure every question and definition is tightly aligned with the material, encourages deep engagement, and avoids any generic or off-topic items."""

        briefing_prompt = """Prepare a concise briefing document titled <BRIEFING_TITLE> that analyzes the main arguments and evidence in the provided text. Structure your briefing as follows:

1. **Main Themes**  
   - Summarize the overarching ideas without using jargon.  
   - Highlight 3–5 central themes showing how they interconnect.

2. **Key Insights and Evidence**  
   - Identify the most important facts, statistics, or examples from the text.  
   - Use bullet points to distinguish between different types of evidence (e.g., data, case studies, expert opinions).

3. **Model Limitations and Alternatives**  
   - Discuss why existing theories or frameworks may fall short.  
   - Reference proposed alternatives and note any open challenges.

4. **Practical Considerations**  
   - Offer specific implications for policymakers, practitioners, or stakeholders.  
   - Address potential risks and benefits.

5. **Illustrative Quotes**  
   - Select 3–5 memorable quotes from the text.  
   - For each quote, provide a one-sentence explanation of its relevance.

Adopt a professional yet accessible tone that balances clarity with intellectual rigor."""

        faq_prompt = """Create a detailed FAQ titled <FAQ_TITLE> to address common questions and concerns arising from the text. For each question, supply a concise, informative answer grounded in the source material. Organize the FAQ into these sections:

1. **Overview Concerns**  
   - Questions about the broader context or foundational issues.  
   - Provide balanced answers that reference historical precedents and current challenges.

2. **Proposed Solutions**  
   - Inquire about remedies or strategies suggested by the text.  
   - Summarize the pros and cons of each approach.

3. **Implementation Challenges**  
   - Address potential barriers to putting solutions into practice.  
   - Include social, economic, or technical obstacles.

4. **Individual Action**  
   - Advice for readers on how to apply insights from the text in their own lives or careers.  
   - Emphasize concrete steps and skill-building.

5. **Ethical and Philosophical Questions**  
   - Tackle deeper questions about values, identity, or long-term impacts.  
   - Ensure answers are nuanced and acknowledge uncertainty where appropriate.

Maintain a clear, engaging style that speaks to a broad audience without oversimplifying."""

        timeline_prompt = """You are a text-analysis assistant. Given the input text, produce:

1. **Timeline Extraction**  
   - List every explicit date or year in the text in chronological order.  
   - For each entry, supply:
     - date: the date in ISO format (YYYY-MM-DD) if possible, or as originally written.
     - event: a one-sentence summary of what occurred on that date.

2. **Character Extraction**  
   - Identify all named persons, organizations, or entities.  
   - For each, supply:
     - name: full name as it appears in the text.
     - role: affiliation or function (e.g., CEO of X, historian, organization).
     - description: 1-2 sentences summarizing their actions or statements.

Please format the output as structured data with clear sections for timeline and characters."""

        dialogue_prompt = """Based on the provided text or file, write an informal, audience-engaging dialogue between two hosts. Follow these guidelines:
1.	Opening
o	Begin with: "Hey everyone, welcome back."
o	Introduce the discussion as a "deep dive" into <TOPIC>.
2.	Structure
o	Alternate speakers (Mandy, Felix).
o	Mix short, punchy lines with longer explanatory segments.
o	Insert affirmations ("Right," "Exactly," "Absolutely") to keep momentum.
3.	Tone & Language
o	Use contractions and colloquial phrases ("You know," "I mean").
o	Keep it energetic and approachable.
o	Use rhetorical questions ("Isn't that wild?") for transitions.
4.	Content Flow
o	Early on, name the source material (articles, studies).
o	Use analogies ("It's like…") to clarify complex ideas.
o	Break points into numbered or clearly signposted segments.
5.	Interaction
o	One host asks questions or expresses confusion; the other responds.
o	Validate each other with phrases like "You've hit the nail on the head."
o	Build collaboratively on each point.
6.	Audience Engagement
o	Directly address listeners ("So to everyone tuning in…").
o	Pose thought-provoking questions for reflection.
7.	Conclusion
o	Signal closing: "So as we wrap things up…"
o	Offer a final takeaway or question.
o	End with: "And on that note…" and a consistent sign-off:
"Until next time, keep <VERB>."
Ensure the dialogue balances informative depth with a casual, friendly vibe."""
        
        return {
            "学习指南": study_guide_prompt,
            "简报文件": briefing_prompt,
            "FAQ文档": faq_prompt,
            "时间线": timeline_prompt,
            "对话": dialogue_prompt
        }
    
    def extract_text_from_file(self, file_path: str) -> str:
        """从文件提取文本"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.processor.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.processor.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.processor.extract_text_from_txt(file_path)
        else:
            return "不支持的文件格式"
    
    def create_output_document(self, results: dict, filename: str) -> str:
        """创建输出Word文档"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(f'文档分析报告: {filename}', 0)
        title.alignment = 1  # 居中
        
        # 添加生成时间
        doc.add_paragraph(f'生成时间: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('') # 空行
        
        # 添加各项分析结果
        for task_name, task_results in results.items():
            # 添加任务标题
            doc.add_heading(task_name, level=1)
            
            if len(task_results) == 1:
                # 单个结果直接添加
                doc.add_paragraph(task_results[0])
            else:
                # 多个结果分别添加
                for i, result in enumerate(task_results, 1):
                    doc.add_heading(f'第 {i} 部分', level=2)
                    doc.add_paragraph(result)
            
            # 添加分隔符
            doc.add_paragraph('').add_run('─' * 50)
            doc.add_paragraph('')  # 空行
        
        # 保存到临时文件
        temp_dir = tempfile.gettempdir()
        output_filename = f"analysis_report_{int(time.time())}.docx"
        output_path = os.path.join(temp_dir, output_filename)
        
        doc.save(output_path)
        return output_path
    
    def get_combined_prompt(self, task_name: str, thinking_mode: str) -> str:
        """组合基础提示词和思维模式前缀"""
        base_prompt = self.base_prompts.get(task_name, "")
        thinking_prefix = self.thinking_modes.get(thinking_mode, "")
        
        if thinking_prefix:
            return f"{thinking_prefix}\n\n{base_prompt}"
        else:
            return base_prompt
    
    def analyze_single_task(self, file_path: str, task_name: str, thinking_mode: str = "标准模式", progress_callback=None) -> dict:
        """分析单个任务"""
        # 提取文本
        if progress_callback:
            progress_callback("正在提取文本...")
        
        text = self.extract_text_from_file(file_path)
        if not text or text == "不支持的文件格式":
            return {"error": "无法提取文本或不支持的文件格式"}
        
        # 分割文本
        text_chunks = self.splitter.split_text(text)
        
        if task_name not in self.base_prompts:
            return {"error": f"未找到任务: {task_name}"}
        
        # 获取组合后的提示词
        prompt = self.get_combined_prompt(task_name, thinking_mode)
        task_results = []
        
        for chunk_idx, chunk in enumerate(text_chunks):
            if progress_callback:
                progress_callback(f"正在处理: {task_name} ({thinking_mode}) - 第 {chunk_idx+1}/{len(text_chunks)} 部分")
            
            chunk_result = ""
            for response_part in self.ollama.generate_stream(prompt, chunk):
                chunk_result += response_part
                # 实时更新进度
                if progress_callback:
                    progress_callback(f"{task_name} ({thinking_mode}) - 第 {chunk_idx+1} 部分: {chunk_result[-50:]}")
            
            task_results.append(chunk_result)
        
        return {f"{task_name} ({thinking_mode})": task_results}

    def analyze_document(self, file_path: str, thinking_mode: str = "标准模式", progress_callback=None) -> dict:
        """分析文档"""
        # 提取文本
        if progress_callback:
            progress_callback("正在提取文本...")
        
        text = self.extract_text_from_file(file_path)
        if not text or text == "不支持的文件格式":
            return {"error": "无法提取文本或不支持的文件格式"}
        
        # 分割文本
        text_chunks = self.splitter.split_text(text)
        
        results = {}
        total_tasks = len(self.base_prompts)
        
        for i, task_name in enumerate(self.base_prompts.keys()):
            if progress_callback:
                progress_callback(f"正在处理: {task_name} ({thinking_mode}) ({i+1}/{total_tasks})")
            
            # 获取组合后的提示词
            prompt = self.get_combined_prompt(task_name, thinking_mode)
            task_results = []
            
            for chunk_idx, chunk in enumerate(text_chunks):
                if progress_callback:
                    progress_callback(f"正在处理: {task_name} ({thinking_mode}) - 第 {chunk_idx+1}/{len(text_chunks)} 部分")
                
                chunk_result = ""
                for response_part in self.ollama.generate_stream(prompt, chunk):
                    chunk_result += response_part
                    # 实时更新进度
                    if progress_callback:
                        progress_callback(f"{task_name} ({thinking_mode}) - 第 {chunk_idx+1} 部分: {chunk_result[-50:]}")
                
                task_results.append(chunk_result)
            
            results[f"{task_name} ({thinking_mode})"] = task_results
        
        return results

def create_course_introduction_interface():
    """创建课程说明界面"""
    
    with gr.Blocks() as interface:
        # 课程标题
        gr.Markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; margin: 20px 0; text-align: center;">
        <h1 style="color: white; font-size: 36px; margin: 0; font-weight: bold;">🎓 提示词工程与人机协作教学</h1>
        <p style="color: #E0E7FF; font-size: 18px; margin: 10px 0;">Prompt Engineering & Human-AI Collaboration</p>
        <p style="color: #C7D2FE; font-size: 16px; margin: 0;">面向人工智能时代的创新教学实践</p>
        </div>
        """)
        
        # 课程概述
        gr.Markdown("""
        ## <span style="color: #1E40AF; font-size: 28px; font-weight: bold;">📚 课程概述</span>
        
        <div style="background-color: #EFF6FF; padding: 20px; border-radius: 12px; margin: 15px 0; border-left: 5px solid #3B82F6;">
        <p style="font-size: 16px; line-height: 1.8; margin: 0;">
        本课程旨在培养学生在人工智能时代的<strong>人机协作能力</strong>，通过系统学习提示词工程技术，
        掌握与大语言模型高效交互的方法，建立相关领域的知识谱系，提升学习和工作效率。
        课程融合了最前沿的AI技术实践，为学生在智能化时代的发展奠定坚实基础。
        </p>
        </div>
        """)
        
        # 授课理念
        gr.Markdown("""
        ---
        ## <span style="color: #7C3AED; font-size: 28px; font-weight: bold;">🎯 授课理念：人机协作，以人为主</span>
        """)
        
        # 人机协作模型可视化
        gr.Markdown("""
        <div style="background: linear-gradient(135deg, #F3F4F6 0%, #E5E7EB 100%); padding: 25px; border-radius: 15px; margin: 20px 0;">
        
        ### <span style="color: #DC2626; font-size: 22px;">🤖 机器的优势：形式系统处理专家</span>
        
        <div style="display: flex; gap: 15px; margin: 20px 0; flex-wrap: wrap;">
            <div style="flex: 1; min-width: 250px; background-color: #FEF2F2; padding: 20px; border-radius: 10px; border: 2px solid #F87171;">
                <h4 style="color: #DC2626; margin-top: 0;">⚡ 符号处理高效</h4>
                <ul style="margin: 10px 0; padding-left: 20px;">
                    <li>快速文本生成与编辑</li>
                    <li>大规模数据处理</li>
                    <li>模式识别与匹配</li>
                </ul>
            </div>
            <div style="flex: 1; min-width: 250px; background-color: #FEF2F2; padding: 20px; border-radius: 10px; border: 2px solid #F87171;">
                <h4 style="color: #DC2626; margin-top: 0;">🔄 媒体转换能力</h4>
                <ul style="margin: 10px 0; padding-left: 20px;">
                    <li>文本 ↔ 图像</li>
                    <li>语音 ↔ 文字</li>
                    <li>结构化数据处理</li>
                </ul>
            </div>
        </div>
        
        ### <span style="color: #059669; font-size: 22px;">🧠 人类的优势：系统外思维者</span>
        
        <div style="display: flex; gap: 15px; margin: 20px 0; flex-wrap: wrap;">
            <div style="flex: 1; min-width: 200px; background-color: #ECFDF5; padding: 20px; border-radius: 10px; border: 2px solid #34D399;">
                <h4 style="color: #059669; margin-top: 0;">👁️ 观察力</h4>
                <p style="margin: 5px 0;">敏锐洞察问题本质<br>发现隐藏的规律</p>
            </div>
            <div style="flex: 1; min-width: 200px; background-color: #ECFDF5; padding: 20px; border-radius: 10px; border: 2px solid #34D399;">
                <h4 style="color: #059669; margin-top: 0;">🤔 反思力</h4>
                <p style="margin: 5px 0;">批判性思维<br>元认知能力</p>
            </div>
            <div style="flex: 1; min-width: 200px; background-color: #ECFDF5; padding: 20px; border-radius: 10px; border: 2px solid #34D399;">
                <h4 style="color: #059669; margin-top: 0;">⚡ 行动力</h4>
                <p style="margin: 5px 0;">决策与执行<br>创新与实践</p>
            </div>
        </div>
        
        ### <span style="color: #7C3AED; font-size: 22px;">🤝 协作模式：1+1>2</span>
        
        <div style="background: linear-gradient(135deg, #8B5CF6 0%, #A78BFA 100%); padding: 20px; border-radius: 12px; color: white; margin: 15px 0;">
            <div style="text-align: center;">
                <h3 style="margin: 0 0 15px 0;">人机协同工作流</h3>
                <div style="font-size: 18px; line-height: 2;">
                    🧠 <strong>人类思考</strong> → 🎯 <strong>提示设计</strong> → 🤖 <strong>AI处理</strong> → 🔍 <strong>人类审查</strong> → ✨ <strong>迭代优化</strong>
                </div>
                <p style="margin: 15px 0 0 0; font-size: 16px; opacity: 0.9;">
                    以人为主导，充分发挥机器的处理能力，实现效率最大化
                </p>
            </div>
        </div>
        
        </div>
        """)
        
        # 前沿技术
        gr.Markdown("""
        ---
        ## <span style="color: #DC2626; font-size: 28px; font-weight: bold;">🚀 前沿技术：掌握AI发展脉搏</span>
        
        <div style="background-color: #FEF2F2; padding: 25px; border-radius: 15px; margin: 20px 0;">
        
        ### <span style="color: #EA580C; font-size: 24px;">🔍 RAG：检索增强生成</span>
        <div style="background-color: #FFF7ED; padding: 20px; border-radius: 10px; margin: 15px 0; border-left: 4px solid #EA580C;">
        <p style="font-size: 16px; margin: 0; line-height: 1.6;">
        <strong>技术核心：</strong> 将外部知识库与大语言模型结合，实现基于事实的精准生成<br>
        <strong>应用价值：</strong> 解决模型幻觉问题，提供可追溯的知识来源<br>
        <strong>实践平台：</strong> 本课程集成RAG文档分析功能，支持PDF/DOCX/TXT智能处理
        </p>
        </div>
        
        ### <span style="color: #7C3AED; font-size: 24px;">🤖 Agent：智能代理系统</span>
        <div style="background-color: #F3F4F6; padding: 20px; border-radius: 10px; margin: 15px 0; border-left: 4px solid #7C3AED;">
        <p style="font-size: 16px; margin: 0; line-height: 1.6;">
        <strong>技术核心：</strong> 具备推理、规划、工具调用能力的自主智能体<br>
        <strong>发展趋势：</strong> 从单一模型向多模态、多工具集成的方向发展<br>
        <strong>教学意义：</strong> 培养学生设计和协作智能代理的能力
        </p>
        </div>
        
        </div>
        """)
        
        # 技术栈
        gr.Markdown("""
        ---
        ## <span style="color: #059669; font-size: 28px; font-weight: bold;">🛠️ 课程技术栈：教学研究一体化平台</span>
        """)
        
        # 技术栈可视化
        gr.Markdown("""
        <div style="background: linear-gradient(135deg, #ECFDF5 0%, #D1FAE5 100%); padding: 25px; border-radius: 15px; margin: 20px 0;">
        
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin: 20px 0;">
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 2px solid #10B981;">
            <h3 style="color: #059669; margin-top: 0; display: flex; align-items: center;">
                <span style="background-color: #10B981; color: white; padding: 5px 10px; border-radius: 20px; margin-right: 10px; font-size: 12px;">核心</span>
                微软 GraphRAG
            </h3>
            <p style="margin: 10px 0; color: #374151;">知识图谱构建与推理<br>企业级RAG解决方案</p>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 2px solid #3B82F6;">
            <h3 style="color: #1E40AF; margin-top: 0; display: flex; align-items: center;">
                <span style="background-color: #3B82F6; color: white; padding: 5px 10px; border-radius: 20px; margin-right: 10px; font-size: 12px;">开源</span>
                Local Deepseek Research
            </h3>
            <p style="margin: 10px 0; color: #374151;">Github开源研究项目<br>本地化部署与优化</p>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 2px solid #8B5CF6;">
            <h3 style="color: #7C3AED; margin-top: 0; display: flex; align-items: center;">
                <span style="background-color: #8B5CF6; color: white; padding: 5px 10px; border-radius: 20px; margin-right: 10px; font-size: 12px;">企业</span>
                腾讯 IMA Deepseek
            </h3>
            <p style="margin: 10px 0; color: #374151;">企业级知识库系统<br>智能问答与检索</p>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 2px solid #F59E0B;">
            <h3 style="color: #D97706; margin-top: 0; display: flex; align-items: center;">
                <span style="background-color: #F59E0B; color: white; padding: 5px 10px; border-radius: 20px; margin-right: 10px; font-size: 12px;">前沿</span>
                ChatGPT 深度研究
            </h3>
            <p style="margin: 10px 0; color: #374151;">提示词工程最佳实践<br>模型能力边界探索</p>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 2px solid #EF4444;">
            <h3 style="color: #DC2626; margin-top: 0; display: flex; align-items: center;">
                <span style="background-color: #EF4444; color: white; padding: 5px 10px; border-radius: 20px; margin-right: 10px; font-size: 12px;">本地</span>
                Ollama 平台
            </h3>
            <p style="margin: 10px 0; color: #374151;">本地大模型部署<br>多模型对话与比较</p>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border: 2px solid #06B6D4;">
            <h3 style="color: #0891B2; margin-top: 0; display: flex; align-items: center;">
                <span style="background-color: #06B6D4; color: white; padding: 5px 10px; border-radius: 20px; margin-right: 10px; font-size: 12px;">开发</span>
                Cursor 平台
            </h3>
            <p style="margin: 10px 0; color: #374151;">AI辅助编程环境<br>智能代码生成与优化</p>
        </div>
        
        </div>
        
        </div>
        """)
        
        # 提示词写作的重要性
        gr.Markdown("""
        ---
        ## <span style="color: #7C2D12; font-size: 28px; font-weight: bold;">🗝️ 提示词工程：AI时代的核心技能</span>
        
        <div style="background: linear-gradient(135deg, #FEF7FF 0%, #F3E8FF 100%); padding: 25px; border-radius: 15px; margin: 20px 0;">
        
        ### <span style="color: #BE185D; font-size: 22px;">💡 技能定位：掌握AI的钥匙</span>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; margin: 15px 0; border-left: 5px solid #BE185D;">
        <p style="font-size: 16px; line-height: 1.8; margin: 0;">
        提示词工程不是过时的技术，而是<strong>人工智能时代的核心能力</strong>。
        正如编程语言是与计算机对话的工具，提示词是与AI对话的语言。
        掌握提示词工程，就是掌握了驾驭AI的能力。
        </p>
        </div>
        
        ### <span style="color: #7C3AED; font-size: 22px;">🧠 学习路径：四个层次递进</span>
        
        <div style="margin: 20px 0;">
        <div style="display: flex; align-items: center; margin: 15px 0; padding: 15px; background-color: #EDE9FE; border-radius: 10px;">
            <div style="background-color: #8B5CF6; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 15px; font-weight: bold;">1</div>
            <div>
                <h4 style="margin: 0; color: #7C3AED;">理解大模型"思维"过程</h4>
                <p style="margin: 5px 0; color: #6B7280;">掌握Transformer架构、注意力机制、推理链等核心概念</p>
            </div>
        </div>
        
        <div style="display: flex; align-items: center; margin: 15px 0; padding: 15px; background-color: #DBEAFE; border-radius: 10px;">
            <div style="background-color: #3B82F6; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 15px; font-weight: bold;">2</div>
            <div>
                <h4 style="margin: 0; color: #1E40AF;">掌握提示词框架</h4>
                <p style="margin: 5px 0; color: #6B7280;">学习CoT、ToT、GoT、CREATES等成熟框架</p>
            </div>
        </div>
        
        <div style="display: flex; align-items: center; margin: 15px 0; padding: 15px; background-color: #D1FAE5; border-radius: 10px;">
            <div style="background-color: #10B981; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 15px; font-weight: bold;">3</div>
            <div>
                <h4 style="margin: 0; color: #059669;">掌握领域知识</h4>
                <p style="margin: 5px 0; color: #6B7280;">结合具体应用场景，建立知识谱系</p>
            </div>
        </div>
        
        <div style="display: flex; align-items: center; margin: 15px 0; padding: 15px; background-color: #FEE2E2; border-radius: 10px;">
            <div style="background-color: #EF4444; color: white; width: 40px; height: 40px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 15px; font-weight: bold;">4</div>
            <div>
                <h4 style="margin: 0; color: #DC2626;">大量实践与迭代</h4>
                <p style="margin: 5px 0; color: #6B7280;">通过不断的输入、测试、改进形成专业技能</p>
            </div>
        </div>
        </div>
        
        </div>
        """)
        
        # 课程特色
        gr.Markdown("""
        ---
        ## <span style="color: #0891B2; font-size: 28px; font-weight: bold;">✨ 课程特色与创新</span>
        
        <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); padding: 25px; border-radius: 15px; margin: 20px 0;">
        
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin: 20px 0;">
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #0891B2; margin-top: 0;">📚 理论与实践并重</h3>
            <ul style="margin: 10px 0; padding-left: 20px; color: #374151;">
                <li>基础知识系统学习</li>
                <li>进阶技术深度探索</li>
                <li>实际项目动手实践</li>
            </ul>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #0891B2; margin-top: 0;">🛠️ 工具平台丰富</h3>
            <ul style="margin: 10px 0; padding-left: 20px; color: #374151;">
                <li>多模型对话体验</li>
                <li>RAG文档分析实战</li>
                <li>提示词优化工具</li>
            </ul>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #0891B2; margin-top: 0;">🎯 面向实际应用</h3>
            <ul style="margin: 10px 0; padding-left: 20px; color: #374151;">
                <li>学术研究场景</li>
                <li>商业应用案例</li>
                <li>创新创业项目</li>
            </ul>
        </div>
        
        <div style="background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #0891B2; margin-top: 0;">🔄 迭代式学习</h3>
            <ul style="margin: 10px 0; padding-left: 20px; color: #374151;">
                <li>从基础到进阶</li>
                <li>理论指导实践</li>
                <li>实践验证理论</li>
            </ul>
        </div>
        
        </div>
        
        </div>
        """)
        
        # 学习成果
        gr.Markdown("""
        ---
        ## <span style="color: #16537e; font-size: 28px; font-weight: bold;">🎖️ 预期学习成果</span>
        
        <div style="background: linear-gradient(135deg, #1E40AF 0%, #3B82F6 100%); padding: 25px; border-radius: 15px; margin: 20px 0; color: white;">
        
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin: 20px 0;">
        
        <div style="background-color: rgba(255,255,255,0.1); padding: 20px; border-radius: 10px;">
            <h3 style="margin-top: 0;">🧠 认知能力提升</h3>
            <ul style="margin: 10px 0; padding-left: 20px;">
                <li>理解AI工作原理</li>
                <li>掌握人机协作模式</li>
                <li>建立系统性思维</li>
            </ul>
        </div>
        
        <div style="background-color: rgba(255,255,255,0.1); padding: 20px; border-radius: 10px;">
            <h3 style="margin-top: 0;">🛠️ 技能能力获得</h3>
            <ul style="margin: 10px 0; padding-left: 20px;">
                <li>熟练使用AI工具</li>
                <li>设计高效提示词</li>
                <li>构建知识管理系统</li>
            </ul>
        </div>
        
        <div style="background-color: rgba(255,255,255,0.1); padding: 20px; border-radius: 10px;">
            <h3 style="margin-top: 0;">💡 创新能力培养</h3>
            <ul style="margin: 10px 0; padding-left: 20px;">
                <li>跨学科整合思维</li>
                <li>问题解决能力</li>
                <li>持续学习适应</li>
            </ul>
        </div>
        
        <div style="background-color: rgba(255,255,255,0.1); padding: 20px; border-radius: 10px;">
            <h3 style="margin-top: 0;">🚀 未来竞争优势</h3>
            <ul style="margin: 10px 0; padding-left: 20px;">
                <li>AI时代核心技能</li>
                <li>数字化工作能力</li>
                <li>终身学习习惯</li>
            </ul>
        </div>
        
        </div>
        
        </div>
        """)
        
        # 结语
        gr.Markdown("""
        ---
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; margin: 20px 0; text-align: center; color: white;">
        <h2 style="margin: 0 0 15px 0; font-size: 24px;">🌟 教学愿景</h2>
        <p style="font-size: 18px; line-height: 1.8; margin: 0; max-width: 800px; margin: 0 auto;">
        通过本课程的学习，学生将能够站在人工智能发展的前沿，
        掌握人机协作的核心技能，在智能化时代中保持学习与创新的主动权，
        成为既懂技术又具备人文思维的复合型人才。
        </p>
        <div style="margin: 20px 0; font-size: 16px; opacity: 0.9;">
        <strong>让我们一起探索AI时代的无限可能！</strong>
        </div>
        </div>
        """)
    
    return interface

def create_prompt_writing_interface():
    """创建提示词写作界面"""
    enhancer = PromptEnhancer()
    
    def enhance_prompt_func(original_prompt, method, progress=gr.Progress()):
        if not original_prompt.strip():
            return "请输入要优化的提示词"
        
        progress(0.1, desc=f"正在使用{method}优化提示词...")
        
        try:
            result = ""
            for response_part in enhancer.enhance_prompt(original_prompt, method):
                result += response_part
                progress(0.5, desc=f"正在生成优化结果...")
            
            progress(1.0, desc="优化完成!")
            return result
            
        except Exception as e:
            return f"优化错误: {str(e)}"
    
    with gr.Column() as interface:
        gr.Markdown("## ✍️ 提示词写作优化")
        
        # 概念解释
        gr.Markdown("""
        ### 🧠 思维框架介绍
        
        **🔗 链式思维 (CoT)**: 引导AI逐步推理，"让我一步步思考..."，提高逻辑性
        
        **🌳 树状思维 (ToT)**: 探索多种解决方案，比较优缺点，选择最佳路径
        
        **🕸️ 图状思维 (GoT)**: 分析概念间关联，构建知识网络，系统性思考
        
        **🌍 全维思维 (EoT)**: 多角度全面分析，时间、空间、理论、实践等维度
        
        **⭐ CO-STAR框架**: 背景-目标-风格-语调-受众-回应，结构化组织提示词
        """)
        
        with gr.Row():
            with gr.Column(scale=2):
                original_input = gr.Textbox(
                    label="📝 输入原始提示词",
                    placeholder="请输入您想要优化的提示词...",
                    lines=5
                )
                
                method_choice = gr.Dropdown(
                    choices=list(enhancer.enhancement_methods.keys()),
                    label="🧠 选择优化方法",
                    value="链式思维 (CoT)"
                )
                
                enhance_btn = gr.Button("🚀 优化提示词", variant="primary")
            
            with gr.Column(scale=3):
                enhanced_output = gr.Textbox(
                    label="✨ 优化后的提示词",
                    lines=15,
                    max_lines=20
                )
        
        # 绑定事件
        enhance_btn.click(
            fn=enhance_prompt_func,
            inputs=[original_input, method_choice],
            outputs=[enhanced_output]
        )
        
        # 使用示例
        gr.Markdown("""
        ### 💡 使用示例
        
        **原始提示词**: "帮我写一篇关于AI的文章"
        
        **优化后效果**: 经过框架优化，会变成结构清晰、目标明确、适合您背景的个性化提示词
        
        ### 📋 使用建议
        1. **输入简洁明确**的原始提示词
        2. **选择合适的优化方法**（复杂任务建议用EoT或GoT）
        3. **复制使用**优化后的提示词与AI对话
        4. **对比效果**，体验不同框架的差异
        """)
    
    return interface

def create_rag_interface():
    """创建RAG文档分析界面"""
    analyzer = DocumentAnalyzer()
    
    def process_single_task(file, task_name, thinking_mode, progress=gr.Progress()):
        if file is None:
            return "请上传文件", "", None
        
        try:
            def update_progress(message):
                progress(0.1, desc=message)
            
            # 分析单个任务
            results = analyzer.analyze_single_task(file.name, task_name, thinking_mode, update_progress)
            
            if "error" in results:
                return results["error"], "", None
            
            # 创建输出文档
            progress(0.9, desc="正在生成输出文档...")
            output_file = analyzer.create_output_document(results, Path(file.name).name)
            
            # 准备显示结果
            result_key = f"{task_name} ({thinking_mode})"
            task_results = results[result_key]
            display_text = f"\n{'='*60}\n{result_key}\n{'='*60}\n"
            
            if len(task_results) == 1:
                display_text += task_results[0]
            else:
                for i, result in enumerate(task_results, 1):
                    display_text += f"\n--- 第 {i} 部分 ---\n"
                    display_text += result
            
            progress(1.0, desc="完成!")
            return "分析完成!", display_text, output_file
            
        except Exception as e:
            return f"处理错误: {str(e)}", "", None
    
    def process_all_tasks(file, thinking_mode, progress=gr.Progress()):
        if file is None:
            return "请上传文件", "", None
        
        try:
            def update_progress(message):
                progress(0.1, desc=message)
            
            # 分析所有任务
            results = analyzer.analyze_document(file.name, thinking_mode, update_progress)
            
            if "error" in results:
                return results["error"], "", None
            
            # 创建输出文档
            progress(0.9, desc="正在生成输出文档...")
            output_file = analyzer.create_output_document(results, Path(file.name).name)
            
            # 准备显示结果
            display_text = ""
            for task_name, task_results in results.items():
                display_text += f"\n{'='*60}\n"
                display_text += f"{task_name}\n"
                display_text += f"{'='*60}\n"
                
                if len(task_results) == 1:
                    display_text += task_results[0]
                else:
                    for i, result in enumerate(task_results, 1):
                        display_text += f"\n--- 第 {i} 部分 ---\n"
                        display_text += result
                
                display_text += "\n\n"
            
            progress(1.0, desc="完成!")
            return "分析完成!", display_text, output_file
            
        except Exception as e:
            return f"处理错误: {str(e)}", "", None
    
    # 创建界面
    with gr.Blocks() as interface:
        gr.Markdown("# 📄 RAG文档智能分析")
        gr.Markdown("上传PDF、DOCX或TXT文件，使用AI进行深度分析并生成学习指南、简报、FAQ等内容")
        
        # RAG概念解释
        with gr.Accordion("💡 什么是RAG？", open=False):
            gr.Markdown("""
            ### 🔍 RAG (Retrieval-Augmented Generation) 概念
            
            **RAG检索增强生成**是一种结合了信息检索和文本生成的AI技术：
            
            - **📚 检索 (Retrieval)**: 从大量文档中找到相关信息
            - **🔗 增强 (Augmented)**: 将检索到的信息作为上下文
            - **✍️ 生成 (Generation)**: 基于检索到的信息生成回答
            
            **优势：**
            - ✅ 基于真实文档内容，减少幻觉
            - ✅ 可以处理最新信息
            - ✅ 提供可追溯的信息来源
            - ✅ 适合知识密集型任务
            
            **本系统的RAG功能：**
            - 📤 上传您的文档
            - 🧠 AI深度理解文档内容  
            - 📋 生成多种格式的分析报告
            - 💾 导出完整的Word文档
            """)
        
        with gr.Row():
            with gr.Column(scale=1):
                file_input = gr.File(
                    label="上传文档",
                    file_types=[".pdf", ".docx", ".txt"],
                    file_count="single"
                )
                
                # 思维模式选择
                thinking_mode = gr.Dropdown(
                    label="🧠 选择思维模式",
                    choices=list(analyzer.thinking_modes.keys()),
                    value="标准模式",
                    info="不同的思维模式会影响AI的分析方式和深度"
                )
                
                gr.Markdown("### 选择分析任务")
                
                # 单独任务按钮
                study_btn = gr.Button("📚 学习指南", variant="secondary", size="sm")
                brief_btn = gr.Button("📊 简报文件", variant="secondary", size="sm")
                faq_btn = gr.Button("❓ FAQ文档", variant="secondary", size="sm")
                timeline_btn = gr.Button("⏰ 时间线", variant="secondary", size="sm")
                dialogue_btn = gr.Button("💬 对话", variant="secondary", size="sm")
                
                # 综合分析按钮
                gr.Markdown("---")
                all_btn = gr.Button("🚀 综合分析（全部任务）", variant="primary", size="lg")
                
                status_output = gr.Textbox(
                    label="状态",
                    placeholder="等待文件上传...",
                    interactive=False
                )
            
            with gr.Column(scale=2):
                result_output = gr.Textbox(
                    label="分析结果",
                    placeholder="分析结果将在这里显示...",
                    lines=20,
                    max_lines=30,
                    interactive=False
                )
        
        download_file = gr.File(label="下载完整报告", interactive=False)
        
        # 绑定事件 - 单独任务
        study_btn.click(
            fn=lambda file, mode, progress=gr.Progress(): process_single_task(file, "学习指南", mode, progress),
            inputs=[file_input, thinking_mode],
            outputs=[status_output, result_output, download_file]
        )
        
        brief_btn.click(
            fn=lambda file, mode, progress=gr.Progress(): process_single_task(file, "简报文件", mode, progress),
            inputs=[file_input, thinking_mode],
            outputs=[status_output, result_output, download_file]
        )
        
        faq_btn.click(
            fn=lambda file, mode, progress=gr.Progress(): process_single_task(file, "FAQ文档", mode, progress),
            inputs=[file_input, thinking_mode],
            outputs=[status_output, result_output, download_file]
        )
        
        timeline_btn.click(
            fn=lambda file, mode, progress=gr.Progress(): process_single_task(file, "时间线", mode, progress),
            inputs=[file_input, thinking_mode],
            outputs=[status_output, result_output, download_file]
        )
        
        dialogue_btn.click(
            fn=lambda file, mode, progress=gr.Progress(): process_single_task(file, "对话", mode, progress),
            inputs=[file_input, thinking_mode],
            outputs=[status_output, result_output, download_file]
        )
        
        # 绑定事件 - 综合分析
        all_btn.click(
            fn=process_all_tasks,
            inputs=[file_input, thinking_mode],
            outputs=[status_output, result_output, download_file]
        )
        
        # 添加说明
        gr.Markdown("""
        ## 📋 使用说明
        
        1. **支持格式**: PDF、DOCX、TXT文件
        2. **思维模式**: 
           - 🎯 **标准模式**: 直接分析，快速高效
           - 🔗 **链式思维 (CoT)**: 逐步推理，逻辑清晰
           - 🌳 **树状思维 (ToT)**: 探索多种可能，择优而行
           - 🕸️ **图状思维 (GoT)**: 关联分析，系统思考
           - 🌍 **全维思维 (EoT)**: 全面分析，多角度审视
        3. **分析模式**: 
           - **单独任务**: 选择特定分析类型，快速完成
           - **综合分析**: 一次性完成所有5种分析
        4. **分析内容**: 
           - 📚 学习指南（总结、理解问题、分析问题、术语表）
           - 📊 简报文件（主要主题、关键见解、实用建议）
           - ❓ FAQ文档（常见问题解答）
           - ⏰ 时间线（重要日期和人物）
           - 💬 对话格式（播客风格讨论）
        
        ## ⚙️ 技术要求
        - 确保Ollama服务运行在 `localhost:11434`
        - 需要安装 `gemma3:4b` 模型: `ollama pull gemma3:4b`
        
        ## 🧪 实验建议
        对同一文档尝试不同思维模式，比较分析质量和深度的差异！
        """)
    
    return interface

def create_knowledge_base_interface():
    """创建基础知识界面"""
    
    with gr.Blocks() as interface:
        gr.Markdown("# 📚 思维框架基础知识")
        gr.Markdown("深入了解各种思维框架的原理、应用方法和实际案例")
        
        # Chain-of-Thought (CoT) 部分
        gr.Markdown("""
        ## <span style="color: #2E86AB; font-size: 24px;">🔗 1. Chain-of-Thought (CoT)</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">📖 定义</span>
        <span style="font-size: 16px;">CoT 要求模型<span style="color: #F18F01; font-weight: bold;">"逐步思考"</span>，在最终回答前先把推理过程拆解为一系列中间步骤。</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">✍️ 提示写法</span>
        在 prompt 中加入：
        - <span style="background-color: #E8F4FD; padding: 5px; border-radius: 5px; font-family: monospace;">"Let's think step by step."</span>
        - <span style="background-color: #E8F4FD; padding: 5px; border-radius: 5px; font-family: monospace;">"请逐步分析，然后给出结论。"</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">📝 写作示例（议论文大纲）</span>
        **任务：** 请为"网络学习的优缺点"写一份 150 词的议论文大纲。
        
        **思维链（CoT）：**
        1. <span style="color: #2E86AB;">首先列出网络学习的三个优点</span>
        2. <span style="color: #2E86AB;">再列出三个缺点</span>
        3. <span style="color: #2E86AB;">对比利弊并得出结论</span>
        4. <span style="color: #F18F01; font-weight: bold;">最后给出完整大纲</span>
        
        **效果：** 模型会先分别罗列正反两方面要点，再组织成清晰的三段式大纲，逻辑更严谨。
        
        ### <span style="color: #A23B72; font-size: 18px;">📄 摘要示例（学术论文）</span>
        **任务：** 请对以下论文段落生成 50 词摘要。
        
        **思维链（CoT）：**
        - **第一步：** 提取研究背景与目的
        - **第二步：** 提取方法与数据
        - **第三步：** 提炼主要结论
        - **最后：** 将前三步内容凝练为一句话
        
        **效果：** 模型会按照指令分步抽取信息，再合并成高度浓缩的摘要。
        
        ---
        """)
        
        # Tree-of-Thought (ToT) 部分
        gr.Markdown("""
        ## <span style="color: #C73E1D; font-size: 24px;">🌳 2. Tree-of-Thought (ToT)</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">📖 定义</span>
        <span style="font-size: 16px;">ToT 在 CoT 的基础上，引入<span style="color: #F18F01; font-weight: bold;">"多条思路分支"</span>——模型生成若干可行思路分支（树的不同分支），并在每一层进行筛选和扩展，最后选出最优路径。</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">✍️ 提示写法</span>
        ```
        请用"树状思维"（Tree-of-Thought）：
        1. 第一层：提出 3 种文章开头思路
        2. 对每个思路打分（给出优缺点）
        3. 在分数最高的思路下再细化 2 种方案
        4. 最终选出最佳方案并写出完整开头段
        ```
        
        ### <span style="color: #A23B72; font-size: 18px;">📝 写作示例（开头段落）</span>
        
        **🌲 第一层生成：**
        1. <span style="color: #2E86AB;">用场景描写引入</span>
        2. <span style="color: #2E86AB;">用数据开篇</span>
        3. <span style="color: #2E86AB;">用提问引出</span>
        
        **🌲 第二层扩展：** 对三种思路各列 2–3 点优劣
        
        **🌲 最终选出：** 假设选"场景描写"并写出 5 句完整开头
        
        ---
        """)
        
        # Graph-of-Thought (GoT) 部分  
        gr.Markdown("""
        ## <span style="color: #7209B7; font-size: 24px;">🕸️ 3. Graph-of-Thought (GoT)</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">📖 定义</span>
        <span style="font-size: 16px;">GoT 将思维节点与节点之间的关联显式化，构建<span style="color: #F18F01; font-weight: bold;">"思维图"</span>（图结构）。每个节点代表一个中间想法，边代表逻辑或因果联系。模型可在图中往返、合并不同思路，更加灵活地综合信息。</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">✍️ 提示写法</span>
        ```
        请用"图状思维"（Graph-of-Thought）：
        - 步骤1：列出 5 个关键观点节点
        - 步骤2：为每对相关节点标注联系（因果、对比、补充）
        - 步骤3：在图中找出最强连接路径，生成文章小标题顺序
        ```
        
        ### <span style="color: #A23B72; font-size: 18px;">📝 写作示例（报告结构）</span>
        
        **🔵 节点设定：**
        1. **节点 A：** 问题描述
        2. **节点 B：** 解决方案一  
        3. **节点 C：** 解决方案二
        4. **节点 D：** 实施计划
        5. **节点 E：** 预期效果
        
        **🔗 关系标注：**
        - **边 AB：** 因果关系
        - **边 BC：** 对比关系  
        - **边 DE：** 逻辑递进
        
        **📋 最终路径：** 问题→方案一→方案二→实施→效果
        
        ---
        """)
        
        # Everything-of-Thought (XoT) 部分
        gr.Markdown("""
        ## <span style="color: #FF6B35; font-size: 24px;">🌍 4. Everything-of-Thought (XoT)</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">📖 定义</span>
        <span style="font-size: 16px;">XoT，或称 "Everything-of-Thought"，是一种<span style="color: #F18F01; font-weight: bold;">全覆盖、全链路的思维策略</span>，它将各种思维框架（CoT、ToT、GoT）、外部知识调用、工具/插件集成、记忆检索与元认知评估等有机融合，旨在让模型在「一步到位」地兼顾所有推理维度与辅助资源后，给出最优解。</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">🎯 核心要素</span>
        
        **1. 🧩 多框架融合**
        - 在单一提示里同时调用 Chain-of-Thought、Tree-of-Thought、Graph-of-Thought 结构，按需灵活组合
        
        **2. 🔌 外部资源接入**  
        - 明确指示模型调用知识库、检索插件、API 接口或自定义函数，以补充实时数据或专业验证
        
        **3. 🧠 记忆与上下文拓展**
        - 利用内置或用户提供的"长期记忆"（之前对话、文档片段）与"短期记忆"（当前上下文）
        
        **4. 🔍 元认知评估**
        - 在最终输出前加入"自我校验"步骤，让模型回顾、评分或对比多个候选结果，择优而出
        
        ### <span style="color: #A23B72; font-size: 18px;">📝 XoT 提示框架示例</span>
        
        **任务：** 为大一新生撰写一份《高效学习方法指南》，约 300 字。
        
        **XoT 提示框架：**
        
        1. **🔗 Chain-of-Thought：**  
           - "请先分步列举四个学习方法要点。"
        
        2. **🌳 Tree-of-Thought：**  
           - "针对每个要点，生成两种不同的展开思路，并简要对比优劣。"
        
        3. **🕸️ Graph-of-Thought：**  
           - "将上述所有要点及其优劣关联成一张思维图（节点＋关联说明），找出最强逻辑路径用于撰写主体段落顺序。"
        
        4. **📚 External Knowledge：**  
           - "引用至少一条教育心理学研究（如'间隔复习'或'刻意练习'）并给出出处。"
        
        5. **💭 Memory Retrieval：**  
           - "结合之前对话中用户提到的'深夜高效复习'经验，融入案例说明。"
        
        6. **🔍 Meta-Cognition Check：**  
           - "最后，给出两个候选版本，并对比哪一个更符合'简洁＋实用'原则，说明理由后选出最佳版本。"
        
        ### <span style="color: #A23B72; font-size: 18px;">🏆 为什么 XoT 有效？</span>
        
        ✅ **一气呵成：** 从"分步思考"到"多路径抉择"，再到"知识引用＋元认知"，一次提示即覆盖所有要点，减少多次反馈。
        
        ✅ **深度与广度兼顾：** 既能精细拆解，也能全局把控，生成结果既具体又有逻辑结构。
        
        ✅ **可验证性：** 外部知识引用与自检步骤，让输出更有说服力、更少偏差。
        
        ---
        """)
        
        # CO-STAR框架补充
        gr.Markdown("""
        ## <span style="color: #16537e; font-size: 24px;">⭐ 5. CO-STAR框架</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">📖 定义</span>
        <span style="font-size: 16px;">CO-STAR 是一个<span style="color: #F18F01; font-weight: bold;">结构化提示词组织框架</span>，通过六个维度系统化地构建高质量提示词。</span>
        
        ### <span style="color: #A23B72; font-size: 18px;">🎯 六大要素</span>
        
        - **📍 C (Context):** 背景信息 - 提供相关上下文
        - **🎯 O (Objective):** 目标任务 - 明确要完成的具体任务  
        - **🎨 S (Style):** 输出风格 - 指定文本风格和语调
        - **💝 T (Tone):** 情感语调 - 设定合适的情感色彩
        - **👥 A (Audience):** 目标受众 - 明确内容的目标读者
        - **📋 R (Response):** 响应格式 - 指定期望的输出结构
        
        ### <span style="color: #A23B72; font-size: 18px;">✨ 使用优势</span>
        
        🎯 **结构清晰** - 六要素确保提示词逻辑完整  
        🎯 **易于调试** - 可以单独优化各个维度  
        🎯 **效果稳定** - 减少模糊性，提高输出一致性  
        🎯 **适用广泛** - 适合各种类型的任务场景
        """)
        
        # 总结部分
        gr.Markdown("""
        ## <span style="color: #2E8B57; font-size: 22px;">📈 思维框架对比总结</span>
        
        | 框架 | 特点 | 适用场景 | 优势 |
        |------|------|----------|------|
        | **🔗 CoT** | 逐步推理 | 逻辑推导、数学计算 | 思路清晰，可追溯 |
        | **🌳 ToT** | 多路径探索 | 创意写作、方案选择 | 选择最优，避免局限 |
        | **🕸️ GoT** | 关系网络 | 复杂分析、系统思考 | 全局视角，关联思维 |
        | **🌍 XoT** | 全维整合 | 综合任务、深度分析 | 全面覆盖，一步到位 |
        | **⭐ CO-STAR** | 结构化 | 提示词构建、标准化 | 格式规范，易于复用 |
        
        ### <span style="color: #A23B72; font-size: 18px;">💡 实用建议</span>
        
        - **🚀 初学者：** 从CoT开始，掌握逐步思考的基本方法
        - **🎨 创意任务：** 使用ToT探索多种可能性
        - **🔬 复杂分析：** 运用GoT构建概念关系网络  
        - **🏆 高端应用：** 尝试XoT进行全维度思考
        - **📝 提示优化：** 使用CO-STAR规范化提示词结构
        """)
    
    return interface

def create_advanced_knowledge_interface():
    """创建进阶知识界面"""
    
    with gr.Blocks() as interface:
        gr.Markdown("# 🎓 提示词工程进阶知识")
        gr.Markdown("深入掌握提示词写作的核心原则、主要框架与高级技术")
        
        # 核心原则部分
        gr.Markdown("""
        ## <span style="color: #1E3A8A; font-size: 28px; font-weight: bold;">🎯 一、核心原则（Underlying Principles）</span>
        
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 15px; border-radius: 10px; margin: 10px 0;">
        <span style="color: white; font-size: 16px; font-weight: bold;">
        在任何具体框架与方法之前，以下几条原则是所有有效提示词写作的基础，必须贯穿始终：
        </span>
        </div>
        """)
        
        # 1. 清晰性与具体性
        gr.Markdown("""
        ### <span style="color: #DC2626; font-size: 22px;">🎯 清晰性与具体性（Clarity & Specificity）</span>
        
        <div style="background-color: #FEF2F2; border-left: 4px solid #DC2626; padding: 15px; margin: 10px 0;">
        <strong style="color: #DC2626;">要点：</strong> 避免模糊、宽泛的指令；对输出内容进行精确描述。<br><br>
        <strong style="color: #DC2626;">原因：</strong> LLM 本质上是概率模型，它会基于提示中出现的关键词与结构进行生成；如果提示过于笼统，模型容易生成偏离预期的结果。
        </div>
        
        **📝 示例对比：**
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #FEE2E2; padding: 15px; border-radius: 8px; border: 2px solid #F87171;">
        <strong style="color: #DC2626;">❌ 不佳示例：</strong><br>
        <span style="font-family: monospace; background-color: #FECACA; padding: 5px; border-radius: 4px;">"写一个故事。"</span>
        </div>
        <div style="flex: 1; background-color: #D1FAE5; padding: 15px; border-radius: 8px; border: 2px solid #34D399;">
        <strong style="color: #059669;">✅ 优良示例：</strong><br>
        <span style="font-family: monospace; background-color: #A7F3D0; padding: 5px; border-radius: 4px;">"写一个关于机器人在反乌托邦未来中学会同情心的短篇故事，整体基调偏向忧郁。"</span>
        </div>
        </div>
        """)
        
        # 2. 上下文至上
        gr.Markdown("""
        ### <span style="color: #7C3AED; font-size: 22px;">👑 上下文至上（Context is King）</span>
        
        <div style="background-color: #F3F4F6; border-left: 4px solid #7C3AED; padding: 15px; margin: 10px 0;">
        <strong style="color: #7C3AED;">要点：</strong> 在提示中预先提供足够的背景信息，明确所指术语与场景，让模型"了解来龙去脉"。<br><br>
        <strong style="color: #7C3AED;">原因：</strong> 模型并不具备真正的常识理解或场景记忆，提示中要主动填补这些信息缺口，减少错误理解的可能性。
        </div>
        
        **🎯 包含内容：**
        - 🔍 定义专业术语
        - 🎭 阐明所处场景与角色  
        - 📋 说明预期输出的用途与格式
        """)
        
        # 3. 角色设定
        gr.Markdown("""
        ### <span style="color: #059669; font-size: 22px;">🎭 角色设定（Role Prompting）</span>
        
        <div style="background-color: #ECFDF5; border-left: 4px solid #059669; padding: 15px; margin: 10px 0;">
        <strong style="color: #059669;">要点：</strong> 为模型分配一个"身份"或"角色"，例如"你是一位资深营销文案策划师"或"你是一位简明扼要的助教"。<br><br>
        <strong style="color: #059669;">原因：</strong> 通过赋予特定角色，能够引导模型在知识面、风格、措辞等方面作出对应调整，从而更贴合目标需求。<br><br>
        <strong style="color: #059669;">效果：</strong> 有助于统一输出口吻，提高专业度和针对性。
        </div>
        """)
        
        # 4. 示例驱动
        gr.Markdown("""
        ### <span style="color: #EA580C; font-size: 22px;">📚 示例驱动（Few-Shot Learning）</span>
        
        <div style="background-color: #FFF7ED; border-left: 4px solid #EA580C; padding: 15px; margin: 10px 0;">
        <strong style="color: #EA580C;">要点：</strong> 在提示中提供若干（通常 1～3 个左右）示例，展示"理想输出"格式与风格。<br><br>
        <strong style="color: #EA580C;">原因：</strong> 比起仅陈述需求，示例更能直观告诉模型如何"去做"，尤其对复杂任务或多步骤任务帮助明显。<br><br>
        <strong style="color: #EA580C;">示例格式：</strong> 若想要生成问答对、摘要或代码片段，都可以在提示里先示范一两个，模型再以此为参照生成新的内容。
        </div>
        """)
        
        # 5. 输出格式控制
        gr.Markdown("""
        ### <span style="color: #0891B2; font-size: 22px;">📐 输出格式控制（Output Format Control）</span>
        
        <div style="background-color: #F0F9FF; border-left: 4px solid #0891B2; padding: 15px; margin: 10px 0;">
        <strong style="color: #0891B2;">要点：</strong> 清楚指明结果应以何种形式呈现，例如"请以要点形式列出"、"请返回 JSON 结构"、"请给出一首押韵的诗歌"等。<br><br>
        <strong style="color: #0891B2;">原因：</strong> 减少后续手动整理、提取等工作量，提升提示的"可用性"。<br><br>
        <strong style="color: #0891B2;">注意：</strong> 若要生成多类型内容（如同时有文字与表格），要在提示里明确区分。
        </div>
        """)
        
        # 6. 参数调节理解
        gr.Markdown("""
        ### <span style="color: #BE185D; font-size: 22px;">⚙️ 参数调节理解（Temperature & Top_P）</span>
        
        <div style="background-color: #FDF2F8; border-left: 4px solid #BE185D; padding: 15px; margin: 10px 0;">
        <strong style="color: #BE185D;">Temperature：</strong> 控制生成内容的随机性，值越高（接近 1），输出越发散；值越低（接近 0），输出越集中于高概率词。<br><br>
        <strong style="color: #BE185D;">Top_P（核采样）：</strong> 控制输出多样性，从概率分布中先筛选出累积概率达到 Top_P 的词汇，再从中随机采样。<br><br>
        <strong style="color: #BE185D;">实践：</strong> 在需要"创造力"或"灵感"类场景时可适当提高 Temperature；在追求"准确、稳定"的场景（如数学推导）时设置较低 Temperature，并可配合 Top_P 低值使用。
        </div>
        """)
        
        # 7. 迭代优化
        gr.Markdown("""
        ### <span style="color: #7C2D12; font-size: 22px;">🔄 迭代优化（Iterative Refinement）</span>
        
        <div style="background-color: #FEF7FF; border-left: 4px solid #7C2D12; padding: 15px; margin: 10px 0;">
        <strong style="color: #7C2D12;">要点：</strong> 绝大多数时候，一次提示难以一次性达到最优效果；要迭代地测试、观察模型输出，并根据输出进行提示微调。
        </div>
        
        **🔄 步骤：**
        1. **🎯 初步设计提示**（基础版）
        2. **🧪 运行模型**，评估输出质量与偏差
        3. **🔧 针对性地改进提示**（增加上下文、修改措辞、添加示例或格式要求）
        4. **🔁 重复以上步骤**，直到满足预期为止
        """)
        
        # 主要框架与方法部分
        gr.Markdown("""
        ---
        ## <span style="color: #1E40AF; font-size: 28px; font-weight: bold;">🏗️ 二、主要框架与方法（Frameworks & Methodologies）</span>
        
        <div style="background: linear-gradient(135deg, #3B82F6 0%, #1E40AF 100%); padding: 15px; border-radius: 10px; margin: 10px 0;">
        <span style="color: white; font-size: 16px; font-weight: bold;">
        以下梳理了当前最常见且已被实践验证的几种提示词写作框架，它们可以单独使用，也可以根据实际需求进行组合或变形。
        </span>
        </div>
        """)
        
        # 1. CREATES 框架
        gr.Markdown("""
        ### <span style="color: #DC2626; font-size: 24px;">🎨 1. CREATES 框架</span>
        
        <div style="background-color: #FEF2F2; padding: 15px; border-radius: 10px; margin: 10px 0;">
        <strong style="color: #DC2626;">来源：</strong> OpenAI 官方提出<br>
        <strong style="color: #DC2626;">定位：</strong> 一个系统化、结构化的提示设计思路，适用于复杂任务<br>
        <strong style="color: #DC2626;">由来：</strong> CREATES 是首字母缩写，对应 6 个关键要素
        </div>
        
        **🔤 六大要素详解：**
        
        | 要素 | 英文 | 说明 | 示例 |
        |------|------|------|------|
        | **📍 C** | Context | 上下文 | 介绍任务应用领域、受众群体、已有数据或前置条件 |
        | **🎭 R** | Role | 角色 | "你是一名经验丰富的中英翻译专家" |
        | **📚 E** | Examples | 示例 | 给出一段示范问答、摘要或代码，让模型"模仿" |
        | **🎯 A** | Task | 任务 | "请将以下英文新闻摘要翻译成中文，并保留原文的关键信息" |
        | **🎵 T** | Tone | 语气 | 正式（formal）、非正式（informal）、幽默（humorous）、客观（objective）等 |
        | **📐 S** | Structure | 结构 | "请以编号要点方式列出，或请以 JSON 格式返回，键名为 'title'、'body' 等" |
        
        **⚖️ 优势与不足：**
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #D1FAE5; padding: 15px; border-radius: 8px;">
        <strong style="color: #059669;">✅ 优势</strong><br>
        • 完整、系统，能够照顾任务的各个环节<br>
        • 对于流程复杂、需要多项要素协同的任务效果显著
        </div>
        <div style="flex: 1; background-color: #FEE2E2; padding: 15px; border-radius: 8px;">
        <strong style="color: #DC2626;">❌ 不足</strong><br>
        • 引入过多要素，提示文字较长<br>
        • 某些要素（例如示例）不易一次性准备，需要额外时间
        </div>
        </div>
        """)
        
        # 2. Chain-of-Thought (CoT)
        gr.Markdown("""
        ### <span style="color: #7C3AED; font-size: 24px;">🔗 2. Chain-of-Thought（CoT）思维链提示</span>
        
        <div style="background-color: #F3F4F6; padding: 15px; border-radius: 10px; margin: 10px 0;">
        <strong style="color: #7C3AED;">核心思想：</strong> 引导模型在给出最终答案之前"写出推理过程"，以提高复杂推理任务的准确率。
        </div>
        
        **🧠 原理：**
        模型在训练过程中学到了一定程度的"隐式推理能力"，通过在提示中加入诸如"让我们一步步分析"或"请解释你的推理过程"，可以激发模型将其内部的推理链"显性化"，从而在做数学运算、逻辑推理、复杂判断时减少遗漏或错误。
        
        **✍️ 典型写法：**
        
        <div style="background-color: #EDE9FE; padding: 15px; border-radius: 8px; margin: 10px 0; font-family: monospace;">
        "问题：23 × 47 等于多少？让我们一步步推理。首先 23×7=161，然后 23×40=920，将二者相加得……最终答案是……"
        </div>
        
        在需要回答推理题时，在提示末尾加上"请详细说明你的思考过程"或"请写出推理步骤，再给出答案"。
        
        **⚖️ 优劣对比：**
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #D1FAE5; padding: 15px; border-radius: 8px;">
        <strong style="color: #059669;">✅ 优势</strong><br>
        • 对解决多步运算题、逻辑题、常识推断题等效果显著<br>
        • 能够生成"可检查的思路"，便于人工核验
        </div>
        <div style="flex: 1; background-color: #FEE2E2; padding: 15px; border-radius: 8px;">
        <strong style="color: #DC2626;">❌ 不足</strong><br>
        • 会显著增加回答长度，不适用于仅需简明答案的场景<br>
        • 在某些任务中可能冗余
        </div>
        </div>
        
        #### <span style="color: #059669; font-size: 20px;">🎯 2.1 零示例思维链（Zero-Shot CoT）</span>
        
        <div style="background-color: #ECFDF5; padding: 15px; border-radius: 8px; margin: 10px 0;">
        <strong style="color: #059669;">概念：</strong> 与 CoT 类似，但不提供示例，仅在提示中加一句"让我们一步步思考"，即可激发模型输出推理过程。<br><br>
        <strong style="color: #059669;">优点：</strong> 实施门槛低，无需准备"示例"<br>
        <strong style="color: #DC2626;">缺点：</strong> 对于特别复杂或高度专业领域的推理任务，零示例版本的效果可能不如完整 CoT
        </div>
        """)
        
        # 3. ReAct 框架
        gr.Markdown("""
        ### <span style="color: #059669; font-size: 24px;">⚡ 3. ReAct（Reason + Act）框架</span>
        
        <div style="background-color: #ECFDF5; padding: 15px; border-radius: 10px; margin: 10px 0;">
        <strong style="color: #059669;">核心思想：</strong> 在"思考（Reason）"与"行动（Act）"之间形成闭环，让模型在相对有限的会话中既能够自我推理，也能主动调用外部工具（如搜索引擎、计算器、API 等），然后再基于获取的信息继续推理与行动。
        </div>
        
        **🔄 具体做法：**
        
        1. **🧠 思考（Reason）：** 首先让模型生成"我需要什么信息""应该使用何种工具"等内部思考
        2. **⚡ 行动（Act）：** 模型将思考结果转化为具体的"工具调用"步骤，例如"搜索：东京最新人口统计数据"，"调用计算器完成 23×47 运算" 等
        3. **🔁 迭代：** 根据工具返回的结果，模型再进行进一步的推理与行动，直到生成最终答案
        
        **📝 典型示例：**
        
        <div style="background-color: #A7F3D0; padding: 15px; border-radius: 8px; margin: 10px 0; font-family: monospace;">
        你是一名研究助理，需要获取东京当前的人口数据。<br>
        1. 思考：我需要使用网络搜索工具来查找最新人口。<br>
        2. 行动：使用搜索引擎查询"2025 年东京人口统计"。<br>
        3. 思考：从返回结果中提取有效数字，并进行验证。<br>
        4. 行动：输出"截至 2025 年，东京人口约为 XXXX 万"。
        </div>
        
        **⚖️ 优势与不足：**
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #D1FAE5; padding: 15px; border-radius: 8px;">
        <strong style="color: #059669;">✅ 优势</strong><br>
        • 适用于需要结合外部信息或多步骤交互的任务<br>
        • 能让 LLM 与现实世界"实时"连接，更加灵活
        </div>
        <div style="flex: 1; background-color: #FEE2E2; padding: 15px; border-radius: 8px;">
        <strong style="color: #DC2626;">❌ 不足</strong><br>
        • 实现复杂，需要先给模型连接到外部工具的能力<br>
        • 对于只需纯文本生成的任务，ReAct 可能显得"多此一举"
        </div>
        </div>
        """)
        
        # 4. 知识生成
        gr.Markdown("""
        ### <span style="color: #EA580C; font-size: 24px;">📚 4. 知识生成（Knowledge Generation, KG）</span>
        
        <div style="background-color: #FFF7ED; padding: 15px; border-radius: 10px; margin: 10px 0;">
        <strong style="color: #EA580C;">核心思想：</strong> 在正式执行任务之前，先让模型"生成与主题相关的知识点"——相当于给模型做一次快速的"预热"，然后再让它基于这些知识展开后续任务。
        </div>
        
        **🔄 操作步骤：**
        
        1. **📖 知识提取：** 提示模型 "请生成关于 [主题] 的一系列要点或事实"
        2. **🎯 任务执行：** 将上述生成的要点作为附加上下文，嵌入到后续的主任务提示中
        3. **📋 输出：** 基于"先生成的知识 + 主任务指令"，生成更具专业性、准确性的回答
        
        **📝 示例：**
        
        <div style="background-color: #FDBA74; padding: 15px; border-radius: 8px; margin: 10px 0; font-family: monospace;">
        第一步提示：请生成一个包含 5 条关于"人工智能在医疗诊断中应用"的事实清单。<br>
        （模型生成：1. … 2. … 3. …）<br><br>
        第二步提示：基于上述 5 条事实，撰写一篇 500 字的分析报告，讨论 AI 技术在医疗诊断领域的机遇与挑战。
        </div>
        
        **⚖️ 优势与不足：**
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #D1FAE5; padding: 15px; border-radius: 8px;">
        <strong style="color: #059669;">✅ 优势</strong><br>
        • 特别适合高度专业或数据密集型场景，让模型提前"进入角色"<br>
        • 提升输出的可靠性与完整性
        </div>
        <div style="flex: 1; background-color: #FEE2E2; padding: 15px; border-radius: 8px;">
        <strong style="color: #DC2626;">❌ 不足</strong><br>
        • 增加了额外一步骤，提示链条更长<br>
        • 若生成的知识要点本身有偏差，则后续任务也会受到影响
        </div>
        </div>
        """)
        
        # 高级技术部分
        gr.Markdown("""
        ---
        ## <span style="color: #7C2D12; font-size: 28px; font-weight: bold;">🚀 三、高级技术（Advanced Techniques）</span>
        
        <div style="background: linear-gradient(135deg, #DC2626 0%, #7C2D12 100%); padding: 15px; border-radius: 10px; margin: 10px 0;">
        <span style="color: white; font-size: 16px; font-weight: bold;">
        当对以下基础方法熟练掌握之后，可以考虑将它们与更高级的提示技术结合，以进一步提升质量或扩展应用场景。
        </span>
        </div>
        """)
        
        # 1. 提示链
        gr.Markdown("""
        ### <span style="color: #1E40AF; font-size: 22px;">🔗 提示链（Prompt Chaining）</span>
        
        <div style="background-color: #EFF6FF; border-left: 4px solid #1E40AF; padding: 15px; margin: 10px 0;">
        <strong style="color: #1E40AF;">定义：</strong> 将一个复杂任务拆分为若干子任务（也可能对应若干子提示），按顺序依次执行，并将每个子任务的输出作为下一个子任务的输入。
        </div>
        
        **📝 示例：**
        - **子提示 1：** 从长篇文章中提取所有关键人物与时间节点
        - **子提示 2：** 根据提取结果，自动生成事件时间线
        - **子提示 3：** 基于事件时间线，撰写一篇结构化的历史分析
        
        **⚖️ 优缺点：**
        - ✅ **优点：** 分而治之，能有效解决单次提示难以覆盖的复杂问题
        - ❌ **缺点：** 流程较长，需要对每步输出结果进行严格校验，否则后续错误会级联
        """)
        
        # 2. 自洽性
        gr.Markdown("""
        ### <span style="color: #7C3AED; font-size: 22px;">🎯 自洽性（Self-Consistency）</span>
        
        <div style="background-color: #F3F4F6; border-left: 4px solid #7C3AED; padding: 15px; margin: 10px 0;">
        <strong style="color: #7C3AED;">定义：</strong> 对同一个提示多次采样生成多个候选答案，然后通过某种方式（例如多数投票或相似度打分）选出"最具一致性"或"最可能正确"的结果。<br><br>
        <strong style="color: #7C3AED;">意义：</strong> 对复杂推理任务，单次生成结果可能存在"走偏"或"遗漏"；通过多次采样，提取具有共性的高置信内容，提高整体准确率。
        </div>
        
        **⚙️ 实施要点：**
        - 需要设置合适的采样参数（如设置较高 Temperature）
        - 考虑后续"筛选"策略，如"关键词匹配度"或"逻辑连贯性"评价
        """)
        
        # 3. 主动提示
        gr.Markdown("""
        ### <span style="color: #059669; font-size: 22px;">❓ 主动提示（Active Prompting）</span>
        
        <div style="background-color: #ECFDF5; border-left: 4px solid #059669; padding: 15px; margin: 10px 0;">
        <strong style="color: #059669;">定义：</strong> 允许模型在正式给出最终答案之前，先对用户的需求进行"澄清性提问"，确保理解无歧义后再执行任务。
        </div>
        
        **📝 示例：**
        
        <div style="background-color: #A7F3D0; padding: 15px; border-radius: 8px; margin: 10px 0; font-family: monospace;">
        用户提示：帮我写一份市场分析报告。<br>
        模型思考（主动提问）：您需要分析的市场是哪个行业？覆盖哪个地区？报告目标受众是谁？<br>
        用户回答后：生成最终报告。
        </div>
        
        **⚖️ 优缺点：**
        - ✅ **优点：** 减少因提示不明确而产生的误差，提高对话效率
        - ❌ **缺点：** 对话回合数增多，用户可能认为过程较"啰嗦"
        """)
        
        # 4. RAG
        gr.Markdown("""
        ### <span style="color: #DC2626; font-size: 22px;">🔍 检索增强生成（Retrieval-Augmented Generation, RAG）</span>
        
        <div style="background-color: #FEF2F2; border-left: 4px solid #DC2626; padding: 15px; margin: 10px 0;">
        <strong style="color: #DC2626;">定义：</strong> 将 LLM 与专门的外部知识库（如向量数据库、文档检索系统等）结合，先检索相关内容，再将检索结果与核心提示一起输入模型进行生成。
        </div>
        
        **🔄 流程：**
        1. **🔍 检索：** 以用户提示为检索查询，找到若干相关文档段落
        2. **🔗 合并：** 将这些文档段落与用户原始提示合并，形成"检索上下文 + 提示"
        3. **📝 生成：** LLM 基于上述扩充上下文生成更精确、带有引用或证据依据的回答
        
        **⚖️ 优势与挑战：**
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #D1FAE5; padding: 15px; border-radius: 8px;">
        <strong style="color: #059669;">✅ 优势</strong><br>
        • 解决 LLM 长尾知识缺失问题<br>
        • 尤其适合需要引用最新数据或特定领域知识的场景
        </div>
        <div style="flex: 1; background-color: #FEE2E2; padding: 15px; border-radius: 8px;">
        <strong style="color: #DC2626;">⚠️ 挑战</strong><br>
        • 需要搭建检索系统、维护索引<br>
        • 并设计好检索与生成之间的衔接策略
        </div>
        </div>
        """)
        
        # 学习资源部分
        gr.Markdown("""
        ---
        ## <span style="color: #0891B2; font-size: 28px; font-weight: bold;">📖 四、进一步学习与资源推荐</span>
        
        ### <span style="color: #DC2626; font-size: 22px;">🌐 在线学习资源</span>
        
        **📚 Learn Prompting（英文）**
        - 网址：https://learnprompting.org/
        - 内容：从基础到进阶的提示教程，并配有大量示例与实践练习
        
        **🔧 OpenAI 文档——提示工程指南（英文）**
        - 网址：https://platform.openai.com/docs/guides/prompt-engineering
        - 内容：官方示例、多种 API 参数解析，以及如何将提示与模型配置结合优化生成质量
        
        **📖 Prompt Engineering Guide（英文）**
        - 网址：https://www.promptingguide.ai/
        - 内容：系统梳理当前主流提示模式与最佳实践；涵盖基础概念、各类框架与高级技术，有社区贡献的示例库
        
        ### <span style="color: #7C3AED; font-size: 22px;">🛠️ 开源项目与示例</span>
        
        **🔗 LangChain（Python 库）**
        - 集成了多种提示模板、提示链与 RAG 示例，适合动手实践
        
        **📝 PromptLibrary**
        - GitHub 上许多社区维护的 prompt 收藏仓库，可直接借鉴真实案例
        """)
        
        # 结论与建议部分
        gr.Markdown("""
        ---
        ## <span style="color: #16537e; font-size: 28px; font-weight: bold;">🎯 结论与建议</span>
        
        ### <span style="color: #DC2626; font-size: 22px;">🎯 结合需求选择框架</span>
        
        <div style="display: flex; gap: 20px; margin: 15px 0;">
        <div style="flex: 1; background-color: #EFF6FF; padding: 15px; border-radius: 8px; border: 2px solid #3B82F6;">
        <strong style="color: #1E40AF;">🚀 简单任务</strong><br>
        可仅依托"清晰与具体""输出格式控制"原则，配合少量示例，即可获得较好效果
        </div>
        <div style="flex: 1; background-color: #FEF2F2; padding: 15px; border-radius: 8px; border: 2px solid #DC2626;">
        <strong style="color: #DC2626;">🔥 复杂任务</strong><br>
        优先考虑 CREATES 全流程或将 CoT/Zero-Shot CoT 与知识生成结合，必要时引入 Prompt Chaining 与 RAG
        </div>
        </div>
        
        ### <span style="color: #7C3AED; font-size: 22px;">🧪 反复实验与迭代</span>
        
        - **🔬 测试优先：** 在正式使用前，先搭建小规模测试，快速评估不同提示组合的效果
        - **🔀 混合使用：** 可将多个提示思路混合：先通过 KG 生成知识，再用 CoT 做过程推理，最后输出
        
        ### <span style="color: #059669; font-size: 22px;">💰 注意成本与效益</span>
        
        - **⏱️ 资源平衡：** 每次模型调用均有延迟与费用，需要在准确性与资源消耗之间平衡
        - **⚡ 优化策略：** 若模型已在特定领域"表现尚可"，可减少示例或上下文长度，以节约额度
        
        ### <span style="color: #EA580C; font-size: 22px;">📈 跟踪前沿动态</span>
        
        <div style="background-color: #FFF7ED; padding: 15px; border-radius: 10px; margin: 10px 0;">
        <strong style="color: #EA580C;">💡 持续学习：</strong> Prompt 工程是一个快速迭代的领域，新方法层出不穷；定期关注社区、学习官方文档与实战范例，有助于发现更高效的策略。
        </div>
        
        ---
        
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 20px; border-radius: 15px; margin: 20px 0; text-align: center;">
        <span style="color: white; font-size: 18px; font-weight: bold;">
        🎓 通过对上述核心原则、常见框架与高级技术的系统梳理，您可以在提示设计之初就建立"全局思路"，进而根据实际场景灵活运用、组合相应方法，不断迭代优化，从而让大语言模型在各类任务中更好地"发挥所长"。
        </span>
        </div>
        """)
    
    return interface

def create_ollama_chat_interface():
    """创建Ollama模型对话界面"""
    
    # 可用模型列表
    available_models = [
        "gemma3:27b",
        "qwen3:32b", 
        "gemma3:12b",
        "deepseek-r1:32b",
        "phi4:latest",
        "openthinker:32b"
    ]
    
    def chat_with_ollama(message, model_name, history, progress=gr.Progress()):
        if not message.strip():
            return history, ""
        
        # 创建临时的Ollama客户端，使用选定的模型
        temp_client = OllamaClient()
        temp_client.model = model_name
        
        try:
            progress(0.1, desc=f"正在使用 {model_name} 处理您的问题...")
            
            # 构建对话历史上下文
            context = ""
            if history:
                for user_msg, bot_msg in history:
                    context += f"用户: {user_msg}\n助手: {bot_msg}\n\n"
            context += f"用户: {message}\n助手: "
            
            # 获取模型响应
            response = ""
            for response_part in temp_client.generate_stream(message):
                response += response_part
                progress(0.5, desc=f"正在生成回答...")
            
            # 更新对话历史
            if history is None:
                history = []
            history.append((message, response))
            
            progress(1.0, desc="完成!")
            return history, ""
            
        except Exception as e:
            error_msg = f"错误: {str(e)}"
            if history is None:
                history = []
            history.append((message, error_msg))
            return history, ""
    
    def clear_history():
        return [], ""
    
    with gr.Blocks() as interface:
        gr.Markdown("# 🤖 Ollama模型对话")
        gr.Markdown("直接与各种Ollama大模型进行对话交流")
        
        # 模型信息展示
        with gr.Accordion("🔍 模型信息", open=False):
            gr.Markdown("""
            ### 🚀 可用模型介绍
            
            | 模型 | 参数量 | 特点 | 适用场景 |
            |------|--------|------|----------|
            | **🔥 gemma3:27b** | 27B | Google最新模型，性能强劲 | 复杂推理、创意写作 |
            | **🌟 qwen3:32b** | 32B | 阿里通义千问，中文优秀 | 中文理解、知识问答 |
            | **⚡ gemma3:12b** | 12B | 平衡性能与速度 | 日常对话、快速响应 |
            | **🧠 deepseek-r1:32b** | 32B | DeepSeek推理模型 | 逻辑推理、数学计算 |
            | **🔬 phi4:latest** | 14B | 微软小参数高性能 | 代码生成、技术问答 |
            | **💭 openthinker:32b** | 32B | 开放思维模型 | 创新思考、头脑风暴 |
            
            ### 💡 使用建议
            - **🚀 高质量任务**: 选择 gemma3:27b 或 qwen3:32b
            - **⚡ 快速响应**: 选择 gemma3:12b 或 phi4:latest  
            - **🧮 逻辑推理**: 选择 deepseek-r1:32b
            - **🎨 创意思考**: 选择 openthinker:32b
            """)
        
        with gr.Row():
            with gr.Column(scale=1):
                # 模型选择
                model_selector = gr.Dropdown(
                    choices=available_models,
                    value="gemma3:12b",
                    label="🤖 选择模型",
                    info="选择要使用的Ollama模型"
                )
                
                # 控制按钮
                clear_btn = gr.Button("🗑️ 清空对话", variant="secondary")
                
                # 模型状态显示
                gr.Markdown("""
                ### 📊 当前会话信息
                - 🌐 服务地址: localhost:11434
                - 🔄 流式输出: 启用
                - 💾 对话记忆: 启用
                
                ### 🎯 使用技巧
                1. **选择合适模型** - 根据任务复杂度选择
                2. **清晰表达** - 具体描述您的需求
                3. **上下文连续** - 利用对话历史功能
                4. **耐心等待** - 大模型需要处理时间
                """)
            
            with gr.Column(scale=2):
                # 对话区域
                chatbot = gr.Chatbot(
                    label="💬 对话区域",
                    height=500,
                    placeholder="在这里显示与AI的对话...",
                    show_label=True
                )
                
                # 输入区域
                with gr.Row():
                    msg_input = gr.Textbox(
                        label="💭 输入消息",
                        placeholder="请输入您想要询问的问题...",
                        lines=2,
                        scale=4
                    )
                    send_btn = gr.Button("📤 发送", variant="primary", scale=1)
        
        # 快捷提示词
        gr.Markdown("### ⚡ 快捷提示词")
        with gr.Row():
            quick_prompts = [
                "请帮我写一份工作总结",
                "解释一下人工智能的发展趋势", 
                "用简单的话解释量子计算",
                "给我一些学习编程的建议"
            ]
            
            for i, prompt in enumerate(quick_prompts):
                quick_btn = gr.Button(f"💡 {prompt}", size="sm")
                quick_btn.click(
                    lambda p=prompt: p,
                    outputs=[msg_input]
                )
        
        # 绑定事件
        msg_input.submit(
            fn=chat_with_ollama,
            inputs=[msg_input, model_selector, chatbot],
            outputs=[chatbot, msg_input]
        )
        
        send_btn.click(
            fn=chat_with_ollama,
            inputs=[msg_input, model_selector, chatbot],
            outputs=[chatbot, msg_input]
        )
        
        clear_btn.click(
            fn=clear_history,
            outputs=[chatbot, msg_input]
        )
        
        # 使用说明
        gr.Markdown("""
        ## 📋 使用说明
        
        1. **🤖 选择模型**: 从下拉菜单中选择适合的模型
        2. **💭 输入问题**: 在文本框中输入您的问题或请求
        3. **📤 发送消息**: 点击发送按钮或按回车键
        4. **💬 查看回答**: 在对话区域查看AI的回答
        5. **🗑️ 清空对话**: 需要时可以清空对话历史
        
        ## ⚙️ 注意事项
        
        - 🔌 确保Ollama服务正在运行 (`ollama serve`)
        - 📥 确保已下载所需模型 (`ollama pull 模型名`)
        - ⏱️ 大模型响应时间较长，请耐心等待
        - 💾 对话历史会在当前会话中保持，切换模型不会清空
        - 🔄 支持上下文理解，可以进行连续对话
        """)
    
    return interface

def create_deep_research_interface():
    """创建深度研究界面"""
    
    def start_deep_research(progress=gr.Progress()):
        """启动深度研究服务"""
        try:
            progress(0.1, desc="正在启动searxng服务...")
            
            # 启动searxng容器
            import subprocess
            result1 = subprocess.run(
                ["docker", "start", "searxng"], 
                capture_output=True, 
                text=True, 
                timeout=30
            )
            
            if result1.returncode != 0:
                return f"启动searxng失败: {result1.stderr}", "❌ 服务启动失败"
            
            progress(0.5, desc="正在启动local-deep-research服务...")
            
            # 启动local-deep-research容器
            result2 = subprocess.run(
                ["docker", "start", "local-deep-research"], 
                capture_output=True, 
                text=True, 
                timeout=30
            )
            
            if result2.returncode != 0:
                return f"启动local-deep-research失败: {result2.stderr}", "❌ 服务启动失败"
            
            progress(1.0, desc="服务启动完成!")
            
            success_msg = """
✅ 深度研究服务启动成功!

🌐 请在浏览器中访问: http://localhost:5000

📊 服务状态:
• searxng: 已启动
• local-deep-research: 已启动

💡 现在您可以开始进行深度研究了！
"""
            return success_msg, "✅ 服务运行中"
            
        except subprocess.TimeoutExpired:
            return "启动超时，请检查Docker服务是否正常运行", "❌ 启动超时"
        except FileNotFoundError:
            return "未找到Docker命令，请确保Docker已正确安装", "❌ Docker未安装"
        except Exception as e:
            return f"启动失败: {str(e)}", "❌ 启动失败"
    
    def stop_deep_research():
        """停止深度研究服务"""
        try:
            import subprocess
            
            # 停止容器
            subprocess.run(["docker", "stop", "searxng"], capture_output=True, timeout=15)
            subprocess.run(["docker", "stop", "local-deep-research"], capture_output=True, timeout=15)
            
            stop_msg = """
🛑 深度研究服务已停止

📊 服务状态:
• searxng: 已停止
• local-deep-research: 已停止

💡 需要重新启动时请点击"启动深度研究"按钮
"""
            return stop_msg, "⏹️ 服务已停止"
            
        except Exception as e:
            return f"停止服务时出错: {str(e)}", "❌ 停止失败"
    
    def check_service_status():
        """检查服务状态"""
        try:
            import subprocess
            
            # 检查容器状态
            result1 = subprocess.run(
                ["docker", "ps", "--filter", "name=searxng", "--format", "{{.Status}}"],
                capture_output=True, text=True
            )
            result2 = subprocess.run(
                ["docker", "ps", "--filter", "name=local-deep-research", "--format", "{{.Status}}"],
                capture_output=True, text=True
            )
            
            searxng_status = "运行中" if result1.stdout.strip() else "已停止"
            research_status = "运行中" if result2.stdout.strip() else "已停止"
            
            status_msg = f"""
📊 当前服务状态:

🔍 searxng: {searxng_status}
🧠 local-deep-research: {research_status}

🌐 访问地址: http://localhost:5000
⏰ 检查时间: {time.strftime("%Y-%m-%d %H:%M:%S")}
"""
            
            overall_status = "✅ 服务运行中" if searxng_status == "运行中" and research_status == "运行中" else "⏸️ 部分或全部服务已停止"
            
            return status_msg, overall_status
            
        except Exception as e:
            return f"检查状态失败: {str(e)}", "❌ 状态检查失败"
    
    with gr.Blocks() as interface:
        # 页面标题
        gr.Markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; margin: 20px 0; text-align: center;">
        <h1 style="color: white; font-size: 36px; margin: 0; font-weight: bold;">🔬 深度研究平台</h1>
        <p style="color: #E0E7FF; font-size: 18px; margin: 10px 0;">Deep Research Platform</p>
        <p style="color: #C7D2FE; font-size: 16px; margin: 0;">基于Docker的本地化深度研究环境</p>
        </div>
        """)
        
        # 平台介绍
        gr.Markdown("""
        ## 🌟 平台概述
        
        <div style="background-color: #EFF6FF; padding: 20px; border-radius: 12px; margin: 15px 0; border-left: 5px solid #3B82F6;">
        <p style="font-size: 16px; line-height: 1.8; margin: 0;">
        深度研究平台集成了<strong>searxng搜索引擎</strong>和<strong>本地化研究工具</strong>，
        为您提供隐私保护的搜索环境和强大的研究分析能力。
        通过Docker容器化部署，确保环境的一致性和安全性。
        </p>
        </div>
        """)
        
        with gr.Row():
            with gr.Column(scale=1):
                # 控制面板
                gr.Markdown("### 🎛️ 服务控制")
                
                start_btn = gr.Button("🚀 启动深度研究", variant="primary", size="lg")
                stop_btn = gr.Button("🛑 停止服务", variant="secondary")
                status_btn = gr.Button("📊 检查状态", variant="secondary")
                
                # 状态显示
                status_indicator = gr.Textbox(
                    label="🔔 服务状态",
                    value="⏸️ 服务未启动",
                    interactive=False
                )
                
                # 快速访问
                gr.Markdown("""
                ### 🌐 快速访问
                
                <div style="background-color: #F0F9FF; padding: 15px; border-radius: 8px; margin: 10px 0;">
                <h4 style="margin-top: 0; color: #1E40AF;">🔗 访问地址</h4>
                <p style="margin: 5px 0;"><strong>主界面:</strong> <a href="http://localhost:5000" target="_blank">http://localhost:5000</a></p>
                <p style="margin: 5px 0; font-size: 14px; color: #6B7280;">点击链接在新窗口中打开深度研究平台</p>
                </div>
                """)
            
            with gr.Column(scale=2):
                # 详细输出
                output_display = gr.Textbox(
                    label="📋 操作日志",
                    lines=12,
                    placeholder="操作结果将在这里显示...",
                    interactive=False
                )
        
        # 绑定事件
        start_btn.click(
            fn=start_deep_research,
            outputs=[output_display, status_indicator]
        )
        
        stop_btn.click(
            fn=stop_deep_research,
            outputs=[output_display, status_indicator]
        )
        
        status_btn.click(
            fn=check_service_status,
            outputs=[output_display, status_indicator]
        )
        
        # 系统要求和说明
        gr.Markdown("""
        ---
        ## 📋 系统要求与说明
        
        ### 🐳 Docker要求
        
        <div style="background-color: #FFF7ED; padding: 20px; border-radius: 12px; margin: 15px 0;">
        
        **必要条件:**
        - ✅ Docker Desktop 已安装并运行
        - ✅ 已拉取以下Docker镜像:
          - `searxng` 容器
          - `local-deep-research` 容器
        
        **启动命令:**
        ```bash
        docker start searxng
        docker start local-deep-research
        ```
        
        </div>
        
        ### 🔧 功能特性
        
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin: 20px 0;">
        
        <div style="background-color: #F0FDF4; padding: 20px; border-radius: 12px; border: 2px solid #22C55E;">
        <h4 style="color: #15803D; margin-top: 0;">🔍 SearXNG搜索</h4>
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li>开源元搜索引擎</li>
        <li>隐私保护搜索</li>
        <li>聚合多个搜索源</li>
        <li>无广告无追踪</li>
        </ul>
        </div>
        
        <div style="background-color: #EFF6FF; padding: 20px; border-radius: 12px; border: 2px solid #3B82F6;">
        <h4 style="color: #1E40AF; margin-top: 0;">🧠 深度研究工具</h4>
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li>本地化研究分析</li>
        <li>数据挖掘与整理</li>
        <li>智能信息提取</li>
        <li>研究报告生成</li>
        </ul>
        </div>
        
        </div>
        
        ### 🚀 使用流程
        
        <div style="background: linear-gradient(135deg, #F3F4F6 0%, #E5E7EB 100%); padding: 20px; border-radius: 12px; margin: 15px 0;">
        
        1. **🔧 环境准备**
           - 确保Docker Desktop运行
           - 确认所需容器已创建
        
        2. **🚀 启动服务**
           - 点击"启动深度研究"按钮
           - 等待两个容器启动完成
        
        3. **🌐 访问平台**
           - 浏览器访问 `http://localhost:5000`
           - 开始您的深度研究之旅
        
        4. **🛑 停止服务**
           - 研究完成后点击"停止服务"
           - 释放系统资源
        
        </div>
        
        ### ⚠️ 注意事项
        
        <div style="background-color: #FEF2F2; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #DC2626;">
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li><strong>端口占用:</strong> 确保5000端口未被其他程序占用</li>
        <li><strong>网络配置:</strong> 确保Docker网络配置正确</li>
        <li><strong>资源消耗:</strong> 深度研究功能会占用一定的CPU和内存资源</li>
        <li><strong>数据安全:</strong> 所有数据在本地处理，确保隐私安全</li>
        </ul>
        </div>
        
        ### 🔧 故障排除
        
        <div style="background-color: #FFFBEB; padding: 15px; border-radius: 8px; margin: 10px 0;">
        <h4 style="color: #D97706; margin-top: 0;">常见问题解决</h4>
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li><strong>容器启动失败:</strong> 检查Docker服务是否运行</li>
        <li><strong>端口被占用:</strong> 使用 <code>netstat -ano | findstr 5000</code> 检查端口</li>
        <li><strong>访问超时:</strong> 等待容器完全启动（约30-60秒）</li>
        <li><strong>页面无法访问:</strong> 确认防火墙设置和网络连接</li>
        </ul>
        </div>
        """)
    
    return interface

def create_graphrag_interface():
    """创建GraphRAG查询界面"""
    
    # 定义虚拟环境和 GraphRag 命令路径
    CONDA_ENV = "graphrag-0.50"
    GRAPH_RAG_COMMAND = "graphrag query"
    ROOT_PATH = "./ragtest"

    def format_response(text):
        """
        Format response text with proper line breaks and separators
        """
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            if any(line.lstrip().startswith(prefix) for prefix in ['1.', '2.', '3.', '•', '-', '※', '*']):
                formatted_lines.append('\n' + line)
            else:
                formatted_lines.append(line)
        
        formatted_text = '\n'.join(formatted_lines)
        
        # 添加更多的章节标记
        sections = ['Summary', 'Conclusion', 'The Threat', 'Humanity\'s Response', 
                   'Internal Conflict', 'The Ending']
        for section in sections:
            if section in formatted_text:
                formatted_text = formatted_text.replace(section, f'\n\n===== {section} =====\n')
        
        return formatted_text.strip()

    def translate_to_chinese(text):
        """
        Translate the refined result to Chinese
        """
        if not text or not isinstance(text, str):
            return "输入文本无效"

        try:
            url = "http://localhost:11434/api/generate"
            headers = {"Content-Type": "application/json"}
            
            prompt = (
                "You are a professional translator. Please translate the following English text to Chinese. "
                "Requirements:\n"
                "1. Maintain the original structure and formatting\n"
                "2. Ensure the translation is natural and fluent in Chinese\n"
                "3. Keep any special formatting, numbers, and section headers\n"
                "4. Use appropriate Chinese punctuation and please provide a high-quality Chinese translation.\n"
                "5. Preserve any technical terms with both English and Chinese translations when necessary\n\n"
                f"Original text:\n{text}\n\n"
                "Please translate:"
            )
            
            data = {
                "model": "gemma3:12b",
                "prompt": prompt,
                "stream": False,
                "temperature": 0.7,
                "max_tokens": 8000,
                "top_p": 0.95
            }
            
            response = requests.post(url, headers=headers, json=data, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                translated = result.get("response", "翻译失败")
                return format_response(translated)
            else:
                return f"请求失败（状态码: {response.status_code}）"
                
        except requests.exceptions.Timeout:
            return "翻译请求超时，请重试"
        except requests.exceptions.ConnectionError:
            return "无法连接到Ollama服务，请确保服务正在运行"
        except Exception as e:
            return f"翻译错误：{str(e)}"

    def refine_result_with_glm4(text):
        """
        Refine and polish query results using Ollama
        """
        if not text or not isinstance(text, str):
            return "输入文本无效"
        
        try:
            url = "http://localhost:11434/api/generate"
            headers = {"Content-Type": "application/json"}
            
            prompt = (
                "Please organize and refine the following text:\n"
                "1. Remove duplicate content and system error messages\n"
                "2. Improve readability and structure\n"
                "3. Maintain the original meaning and key information\n"
                "4. Create a coherent, well-organized narrative\n"
                "5. Highlight main themes and important points\n\n"
                f"Text to process:\n{text}\n\n"
                "Please provide the refined version:"
            )
            
            data = {
                "model": "gemma3:12b",
                "prompt": prompt,
                "stream": False,
                "temperature": 0.7,
                "max_tokens": 8000,
                "top_p": 0.95,
                "frequency_penalty": 0.5
            }
            
            response = requests.post(url, headers=headers, json=data, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                refined = result.get("response", "优化失败")
                return format_response(refined)
            else:
                return f"请求失败（状态码: {response.status_code}）"
                
        except requests.exceptions.Timeout:
            return "优化请求超时，请重试"
        except requests.exceptions.ConnectionError:
            return "无法连接到Ollama服务，请确保服务正在运行"
        except Exception as e:
            return f"优化错误：{str(e)}"

    def graphrag_query(query, method):
        """Execute GraphRag query"""
        if not query.strip():
            return "查询内容不能为空"
        try:
            command = (
                f"conda run -n {CONDA_ENV} {GRAPH_RAG_COMMAND} "
                f"--root {ROOT_PATH} --method {method} --query \"{query}\""
            )
            result = subprocess.check_output(command, shell=True, text=True, stderr=subprocess.STDOUT)
            save_query_result(query, result, method)
            return result
        except subprocess.CalledProcessError as e:
            error_message = f"GraphRAG查询失败:\n错误代码: {e.returncode}\n错误信息: {e.output}"
            save_query_result(query, error_message, method)
            return error_message
        except Exception as e:
            return f"查询过程中发生意外错误: {str(e)}"

    def save_query_result(query, result, method):
        """Save query results to file"""
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = "./query_results"
        os.makedirs(output_dir, exist_ok=True)
        filename = os.path.join(output_dir, f"query_results_{timestamp}_{method}.txt")

        with open(filename, "w", encoding="utf-8") as file:
            file.write(f"Query: {query}\nMethod: {method}\n\nResult:\n{result}")

    # 预设问题集
    PRESET_QUESTIONS = [
        {
            "question": "What are the fundamental frameworks and methodologies for effective prompt engineering?",
            "method": "global",
            "category": "Framework",
            "description": "提示词工程的基础框架和方法论"
        },
        {
            "question": "How do different prompt writing techniques compare in terms of effectiveness and use cases?",
            "method": "global", 
            "category": "Comparison",
            "description": "不同提示词技术的效果对比分析"
        },
        {
            "question": "What are the specific step-by-step processes for constructing high-quality prompts?",
            "method": "local",
            "category": "Process", 
            "description": "构建高质量提示词的具体步骤"
        },
        {
            "question": "Can you provide concrete examples of successful prompt templates for different AI tasks?",
            "method": "local",
            "category": "Examples",
            "description": "不同AI任务的成功提示词模板示例"
        },
        {
            "question": "What are the common pitfalls and mistakes to avoid when writing prompts?",
            "method": "local",
            "category": "Best Practices",
            "description": "提示词写作中的常见误区和避免方法"
        },
        {
            "question": "How has prompt engineering evolved and what are the emerging trends in this field?",
            "method": "drift",
            "category": "Evolution",
            "description": "提示词工程的发展历程和新兴趋势"
        },
        {
            "question": "What role does context length and structure play in prompt effectiveness?",
            "method": "local",
            "category": "Technical",
            "description": "上下文长度和结构对提示词效果的影响"
        },
        {
            "question": "How do domain-specific prompting strategies differ across various industries and applications?",
            "method": "global",
            "category": "Domain-Specific",
            "description": "不同行业和应用的领域特定提示词策略"
        },
        {
            "question": "What are the psychological and cognitive principles behind effective prompt design?",
            "method": "drift",
            "category": "Psychology",
            "description": "有效提示词设计背后的心理学和认知原理"
        },
        {
            "question": "How can prompt writers measure and evaluate the quality and performance of their prompts?",
            "method": "local",
            "category": "Evaluation",
            "description": "提示词质量和性能的衡量评估方法"
        }
    ]

    # 界面构建
    with gr.Blocks() as interface:
        # 页面标题
        gr.Markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; margin: 20px 0; text-align: center;">
        <h1 style="color: white; font-size: 36px; margin: 0; font-weight: bold;">🕸️ GraphRAG查询系统</h1>
        <p style="color: #E0E7FF; font-size: 18px; margin: 10px 0;">Graph-based Retrieval Augmented Generation</p>
        <p style="color: #C7D2FE; font-size: 16px; margin: 0;">基于知识图谱的高级查询与分析平台</p>
        </div>
        """)

        # GraphRAG介绍
        gr.Markdown("""
        ## 🌟 GraphRAG概述
        
        <div style="background-color: #EFF6FF; padding: 20px; border-radius: 12px; margin: 15px 0; border-left: 5px solid #3B82F6;">
        <p style="font-size: 16px; line-height: 1.8; margin: 0;">
        <strong>GraphRAG</strong>是微软开发的下一代检索增强生成技术，通过构建知识图谱来理解文档中的复杂关系。
        相比传统的RAG技术，GraphRAG能够更好地处理<strong>多跳推理</strong>和<strong>全局理解</strong>任务，
        特别适合分析复杂文档、研究报告和大型知识库。
        </p>
        </div>
        """)

        # GraphRAG vs 传统RAG对比
        with gr.Accordion("🔍 GraphRAG vs 传统RAG", open=False):
            gr.Markdown("""
            ### 📊 技术对比

            | 特性 | 传统RAG | GraphRAG |
            |------|---------|----------|
            | **检索方式** | 向量相似度 | 知识图谱 + 向量检索 |
            | **理解能力** | 局部片段理解 | 全局关系理解 |
            | **推理能力** | 单跳检索 | 多跳关系推理 |
            | **适用场景** | 简单问答 | 复杂分析、洞察发现 |
            | **准确性** | 中等 | 高 |
            | **计算成本** | 低 | 中等 |

            ### 🎯 GraphRAG的优势
            - **🕸️ 图谱构建**: 自动从文档中提取实体和关系
            - **🔍 多层检索**: 支持local、global、drift三种查询模式
            - **🧠 深度理解**: 能够理解复杂的语义关系
            - **📈 可扩展性**: 适合大规模文档处理
            """)

        # 预设问题集
        with gr.Accordion("🎯 提示词写作深度挖掘问题集 - 一键查询", open=True):
            gr.Markdown("""
            ### 📋 精选问题列表
            点击下方按钮可以直接执行相应的GraphRAG查询，无需手动输入问题。
            """)
            
            # 按类别分组显示问题
            categories = {}
            for q in PRESET_QUESTIONS:
                if q["category"] not in categories:
                    categories[q["category"]] = []
                categories[q["category"]].append(q)
            
            # 存储所有预设按钮
            preset_buttons = []
            
            for category, questions in categories.items():
                with gr.Group():
                    gr.Markdown(f"#### 🏷️ {category}")
                    for i, q in enumerate(questions):
                        with gr.Row():
                            gr.Markdown(f"**{q['description']}** ({q['method'].upper()})")
                            preset_btn = gr.Button(
                                f"🚀 查询", 
                                variant="secondary", 
                                size="sm"
                            )
                            preset_buttons.append((preset_btn, q))

        # 查询设置区域
        with gr.Row():
            with gr.Column(scale=2):
                query_input = gr.Textbox(
                    label="🔍 查询内容",
                    placeholder="例如: Who is the English patient? 或 What are the main themes in the document?",
                    lines=3,
                    info="支持复杂的分析性问题和多跳推理查询"
                )
                
            with gr.Column(scale=1):
                method_dropdown = gr.Dropdown(
                    choices=["local", "global", "drift"],
                    label="📊 查询方法",
                    value="local",
                    info="选择GraphRAG查询模式，详见下方说明"
                )
                
                # 查询方法简要说明
                gr.Markdown("""
                <div style="background-color: #F8FAFC; padding: 10px; border-radius: 6px; margin: 5px 0; font-size: 13px;">
                <strong>🔍 Local:</strong> 基于本地社区的查询，适合具体问题<br>
                <strong>🌍 Global:</strong> 全数据集查询，适合宏观分析<br>
                <strong>🌊 Drift:</strong> 探索性查询，发现潜在关联
                </div>
                """)
        
        # 查询方法说明
        with gr.Accordion("💡 查询方法说明", open=False):
            gr.Markdown("""
            ### 🎯 三种查询模式

            **🔍 Local 查询模式 (本地社区查询)**
            - **原理**: 基于预构建的本地社区结构进行查询
            - **特点**: 查询范围聚焦于特定的实体群组和局部关系
            - **适用场景**: 具体事实查询、实体间关系、局部信息检索
            - **优势**: 响应速度快，计算成本低，结果精准
            - **示例查询**: 
              - "Who is John Smith and what is his role?"
              - "张三的具体工作职责是什么？"
              - "What happened in Chapter 3?"

            **🌍 Global 查询模式 (全局数据集查询)**
            - **原理**: 基于整个知识图谱进行全局分析和推理
            - **特点**: 能够跨越多个社区，整合全局信息
            - **适用场景**: 主题总结、趋势分析、宏观洞察
            - **优势**: 视野广阔，能发现全局模式和深层关联
            - **示例查询**:
              - "What are the main themes across all documents?"
              - "整个文档集的核心主题和趋势是什么？"
              - "How do different concepts relate globally?"

            **🌊 Drift 查询模式 (概念漂移查询)**
            - **原理**: 探索概念在不同上下文中的语义漂移和演化
            - **特点**: 能够发现概念的动态变化和潜在关联
            - **适用场景**: 探索性研究、概念演化分析、创新发现
            - **优势**: 发现隐藏模式，识别新兴趋势，探索未知关联
            - **示例查询**:
              - "How has the concept of AI evolved over time?"
              - "人工智能概念在不同文档中的演化轨迹？"
              - "What emerging patterns can be identified?"
            """)

        # 状态显示
        status_display = gr.Markdown("🟢 系统就绪，等待查询...", elem_classes=["status"])
        
        # 操作按钮
        with gr.Row():
            query_btn = gr.Button("🚀 执行GraphRAG查询", variant="primary", size="lg")
            refine_btn = gr.Button("✨ 优化结果", variant="secondary")
            translate_btn = gr.Button("🈶 翻译成中文", variant="secondary")
            clear_btn = gr.Button("🗑️ 清空所有", variant="stop")
        
        # 结果显示区域
        with gr.Tabs():
            with gr.TabItem("📄 原始查询结果"):
                raw_result = gr.Textbox(
                    label="GraphRAG原始查询结果",
                    lines=15,
                    max_lines=30,
                    placeholder="查询结果将在这里显示..."
                )
            
            with gr.TabItem("✨ 优化结果"):
                refined_result = gr.Textbox(
                    label="AI优化后的结果",
                    lines=15,
                    max_lines=30,
                    placeholder="优化后的结果将在这里显示..."
                )
            
            with gr.TabItem("🈶 中文翻译"):
                translated_result = gr.Textbox(
                    label="中文翻译结果",
                    lines=15,
                    max_lines=30,
                    placeholder="中文翻译将在这里显示..."
                )

        # 执行查询函数
        def query_action(query, method, progress=gr.Progress()):
            if not query.strip():
                return "⚠️ 查询内容不能为空", gr.update(value="")
            
            try:
                progress(0.1, desc="正在执行GraphRAG查询...")
                result = graphrag_query(query, method)
                progress(1.0, desc="查询完成!")
                return "✅ GraphRAG查询执行完成", result
            except Exception as e:
                return f"❌ 查询出错: {str(e)}", ""
        

        # 优化结果函数
        def refine_action(text, progress=gr.Progress()):
            if not text.strip():
                return "⚠️ 没有可优化的内容", gr.update(value="")
            
            try:
                progress(0.1, desc="正在使用AI优化结果...")
                result = refine_result_with_glm4(text)
                progress(1.0, desc="优化完成!")
                return "✅ 结果优化完成", result
            except Exception as e:
                return f"❌ 优化出错: {str(e)}", ""
        
        # 翻译结果函数
        def translate_action(text, progress=gr.Progress()):
            if not text.strip():
                return "⚠️ 没有可翻译的内容", gr.update(value="")
            
            try:
                progress(0.1, desc="正在翻译成中文...")
                result = translate_to_chinese(text)
                progress(1.0, desc="翻译完成!")
                return "✅ 翻译完成", result
            except Exception as e:
                return f"❌ 翻译出错: {str(e)}", ""
        
        # 清空所有函数
        def clear_action():
            return (
                "🔄 系统已重置，准备新的查询", 
                gr.update(value=""), 
                gr.update(value=""), 
                gr.update(value=""), 
                gr.update(value="")
            )
        
        # 事件绑定
        query_btn.click(
            query_action, 
            inputs=[query_input, method_dropdown], 
            outputs=[status_display, raw_result]
        )
        
        # 为每个预设问题按钮绑定事件
        for btn, question_data in preset_buttons:
            def create_preset_handler(q_data):
                def handler(progress=gr.Progress()):
                    try:
                        progress(0.1, desc=f"正在执行预设查询 ({q_data['method'].upper()})...")
                        result = graphrag_query(q_data["question"], q_data["method"])
                        progress(1.0, desc="查询完成!")
                        return (
                            f"✅ 预设问题查询完成",
                            q_data["question"], 
                            q_data["method"], 
                            result
                        )
                    except Exception as e:
                        return f"❌ 预设查询出错: {str(e)}", "", "local", ""
                return handler
            
            btn.click(
                create_preset_handler(question_data),
                outputs=[status_display, query_input, method_dropdown, raw_result]
            )
        
        refine_btn.click(
            refine_action, 
            inputs=[raw_result], 
            outputs=[status_display, refined_result]
        )
        
        translate_btn.click(
            translate_action, 
            inputs=[raw_result], 
            outputs=[status_display, translated_result]
        )
        
        clear_btn.click(
            clear_action,
            inputs=[],
            outputs=[status_display, query_input, raw_result, refined_result, translated_result]
        )

        # 系统要求和配置
        gr.Markdown("""
        ---
        ## ⚙️ 系统配置与要求
        
        ### 🐍 Conda环境配置
        
        <div style="background-color: #FFF7ED; padding: 20px; border-radius: 12px; margin: 15px 0;">
        
        **必要条件:**
        - ✅ Conda环境名称: `graphrag-0.50`
        - ✅ GraphRAG包已正确安装
        - ✅ 数据目录: `./ragtest`
        - ✅ Ollama服务运行在localhost:11434
        
        **安装命令:**
        ```bash
        conda create -n graphrag-0.50 python=3.10
        conda activate graphrag-0.50
        pip install graphrag==0.50
        ```
        
        </div>
        
        ### 📂 目录结构
        
        <div style="background-color: #F0FDF4; padding: 15px; border-radius: 8px; margin: 10px 0;">
        
        ```
        ./ragtest/
        ├── input/                        # 输入文档目录
        ├── output/                       # GraphRAG处理结果
        │   ├── artifacts/                # 生成的图谱数据和中间文件
        │   ├── create_final_community_reports.parquet
        │   ├── create_final_entities.parquet
        │   ├── create_final_relationships.parquet
        │   └── create_final_text_units.parquet
        ├── settings.yaml                 # GraphRAG配置文件
        └── .env                         # 环境变量配置
        ```
        
        </div>
        
        ### 🔧 使用流程
        
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 15px; margin: 20px 0;">
        
        <div style="background-color: #EFF6FF; padding: 15px; border-radius: 8px; text-align: center;">
        <h4 style="color: #1E40AF; margin-top: 0;">1️⃣ 准备数据</h4>
        <p style="margin: 5px 0;">将文档放入input目录</p>
        </div>
        
        <div style="background-color: #F0FDF4; padding: 15px; border-radius: 8px; text-align: center;">
        <h4 style="color: #059669; margin-top: 0;">2️⃣ 构建图谱</h4>
        <p style="margin: 5px 0;">运行graphrag index命令</p>
        </div>
        
        <div style="background-color: #FEF2F2; padding: 15px; border-radius: 8px; text-align: center;">
        <h4 style="color: #DC2626; margin-top: 0;">3️⃣ 执行查询</h4>
        <p style="margin: 5px 0;">使用本界面进行查询</p>
        </div>
        
        <div style="background-color: #FFFBEB; padding: 15px; border-radius: 8px; text-align: center;">
        <h4 style="color: #D97706; margin-top: 0;">4️⃣ 结果优化</h4>
        <p style="margin: 5px 0;">AI优化和中文翻译</p>
        </div>
        
        </div>
        
        ### 🎯 预设问题集使用说明
        
        <div style="background-color: #F0F9FF; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #0EA5E9;">
        <h4 style="color: #0C4A6E; margin-top: 0;">📋 10个精选深度问题</h4>
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li><strong>Framework:</strong> 提示词工程基础框架和方法论</li>
        <li><strong>Comparison:</strong> 不同技术的效果对比分析</li>
        <li><strong>Process:</strong> 构建高质量提示词的具体步骤</li>
        <li><strong>Examples:</strong> 成功提示词模板和案例</li>
        <li><strong>Best Practices:</strong> 常见误区和最佳实践</li>
        <li><strong>Evolution:</strong> 发展历程和新兴趋势</li>
        <li><strong>Technical:</strong> 技术细节和优化要点</li>
        <li><strong>Domain-Specific:</strong> 领域特定策略</li>
        <li><strong>Psychology:</strong> 心理学和认知原理</li>
        <li><strong>Evaluation:</strong> 质量评估和测量方法</li>
        </ul>
        </div>
        
        ### ⚠️ 注意事项
        
        <div style="background-color: #FEF2F2; padding: 15px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #DC2626;">
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li><strong>环境激活:</strong> 确保conda环境正确激活</li>
        <li><strong>数据准备:</strong> 需要预先建立GraphRAG索引</li>
        <li><strong>计算资源:</strong> GraphRAG查询可能需要较长时间</li>
        <li><strong>模型依赖:</strong> 优化和翻译功能依赖Ollama gemma3:12b模型</li>
        <li><strong>网络连接:</strong> 确保localhost:11434可访问</li>
        </ul>
        </div>
        
        ### 🔧 故障排除
        
        <div style="background-color: #FFFBEB; padding: 15px; border-radius: 8px; margin: 10px 0;">
        <h4 style="color: #D97706; margin-top: 0;">常见问题解决</h4>
        <ul style="margin: 10px 0; padding-left: 20px;">
        <li><strong>环境错误:</strong> 检查conda环境是否正确创建和激活</li>
        <li><strong>路径问题:</strong> 确认./ragtest目录存在且有权限</li>
        <li><strong>查询超时:</strong> 复杂查询可能需要更长时间</li>
        <li><strong>翻译失败:</strong> 检查Ollama服务和gemma3:12b模型</li>
        <li><strong>优化失败:</strong> 确认原始结果不为空且格式正确</li>
        </ul>
        </div>
        
        ---
        
        <div style="text-align: center; margin: 20px 0; color: #6B7280;">
        <strong>© Designed by Toby LUO@2025.01</strong><br>
        GraphRAG集成版本 - 提供专业的知识图谱查询体验
        </div>
        """)
    
    return interface

def create_main_interface():
    """创建主界面"""
    with gr.Blocks(title="提示词写作小课堂", theme=gr.themes.Soft()) as demo:
        gr.Markdown("# 🤖 Toby AI课堂--提示词写作基础")
        gr.Markdown("选择您需要的功能模块，开始AI写作之旅！")
        
        with gr.Tabs() as tabs:
            with gr.TabItem("📖 课程说明"):
                course_interface = create_course_introduction_interface()
                
            with gr.TabItem("📚 基础知识"):
                knowledge_interface = create_knowledge_base_interface()
            
            with gr.TabItem("🎓 进阶知识"):
                advanced_knowledge_interface = create_advanced_knowledge_interface()
            
            with gr.TabItem("✍️ 提示词优化"):
                prompt_interface = create_prompt_writing_interface()
            
            with gr.TabItem("📄 RAG文档分析"):
                rag_interface = create_rag_interface()
                
            with gr.TabItem("🤖 模型实战"):
                ollama_interface = create_ollama_chat_interface()
                
            with gr.TabItem("🔬 深度研究"):
                deep_research_interface = create_deep_research_interface()
                
            with gr.TabItem("🕸️ GraphRAG"):
                graphrag_interface = create_graphrag_interface()
        
        # 全局说明
        gr.Markdown("""
        ---
        ## 🚀 平台功能
        
        ### 📖 课程说明
        - 了解提示词工程与人机协作的核心理念
        - 掌握人机协作的最佳实践模式  
        - 学习前沿AI技术和发展趋势
        - 建立系统化的AI学习思路
        
        ### 📚 基础知识  
        - 深入学习CoT、ToT、GoT、XoT等思维框架
        - 掌握各种框架的应用方法和使用场景
        - 理解框架原理，提升提示词设计能力
        
        ### 🎓 进阶知识
        - 掌握提示词工程的核心原则和最佳实践
        - 学习CREATES、ReAct、知识生成等高级框架
        - 了解提示链、自洽性、RAG等前沿技术
        
        ### ✍️ 提示词写作优化
        - 使用先进的思维框架优化您的提示词
        - 支持CoT、ToT、GoT、EoT和CO-STAR框架
        - 提高AI交互效果和回答质量
        
        ### 📄 RAG文档分析  
        - 上传文档进行智能分析
        - 生成学习指南、简报、FAQ等多种格式
        - 支持多种思维模式深度分析
        - 导出完整的Word分析报告
        
        ### 🤖 模型对话
        - 直接与多种Ollama大模型对话
        - 支持6种不同规模和特色的模型
        - 实时流式对话体验
        - 保持上下文的连续对话
        
        ### 🔬 深度研究
        - 启动本地化深度研究平台
        - 集成SearXNG隐私搜索引擎
        - 提供专业的研究分析工具
        - 支持Docker容器化部署
        
        ### 🕸️ GraphRAG
        - 基于知识图谱的高级查询系统
        - 支持local、global、drift三种查询模式
        - 集成AI结果优化和中文翻译
        - 微软GraphRAG技术的完整实现
        
        ## ⚙️ 系统要求
        - Ollama服务: `http://localhost:11434`
        - 推荐模型: `gemma3:12b` (平衡性能)
        - 启动命令: `ollama serve`
        - 模型下载: `ollama pull 模型名`
        - Docker环境: 深度研究功能需要Docker支持
        - Conda环境: GraphRAG需要graphrag-0.50环境
        
        ---
        ### 📝 版权信息
        **© Toby LUO@2025.06** - Toby AI课堂
        """)
    
    return demo

if __name__ == "__main__":
    # 检查依赖
    required_packages = {
        "gradio": "gradio",
        "requests": "requests", 
        "PyPDF2": "PyPDF2",
        "python-docx": "docx",
        "pdfplumber": "pdfplumber"
    }
    
    missing_packages = []
    for package_name, import_name in required_packages.items():
        try:
            __import__(import_name)
        except ImportError:
            missing_packages.append(package_name)
    
    if missing_packages:
        print("缺少以下依赖包，请安装:")
        for package in missing_packages:
            print(f"  pip install {package}")
        exit(1)
    
    # 启动应用
    demo = create_main_interface()
    demo.launch(
        server_name="127.0.0.1",
        server_port=7860,
        share=True,
        show_error=True
    ) 